import os
from datetime import datetime, timezone
from fastapi import FastAPI, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session, joinedload
from database import engine, get_db, Base, SessionLocal
from fastapi import Request
from models import (
    User, Document, Tag, Approval, Comment, Notification,
    History, Version, Attachment, ApprovalRoute, doc_tags, doc_related, Task, Resolution,
    AuditLog, DocumentTemplate, NomenclatureCase,
)
from schemas import (
    UserRegister, UserLogin, UserOut, UserOutPublic, Token, UserCreate, DeputySet,
    DocumentCreate, DocumentOut, ApprovalOut, CommentOut, CommentCreate,
    ApprovalAction, NotificationOut, RouteCreate, RouteOut, TagOut,
    TaskCreate, TaskUpdate, TaskOut, ResolutionOut, ResolutionCreate,
    ProfileUpdate, PasswordChange,
    AuditLogOut, TemplateCreate, TemplateOut, BulkAction,
    NomenclatureCaseCreate, NomenclatureCaseOut,
)
from auth import hash_password, verify_password, create_token, get_current_user
import secrets
import shutil
import json

import re
import time
from collections import defaultdict

# Rate limiting for login
_login_attempts: dict[str, list[float]] = defaultdict(list)
_LOGIN_MAX_ATTEMPTS = 5
_LOGIN_WINDOW = 300  # 5 minutes
_login_last_cleanup = 0.0

def _check_rate_limit(key: str):
    global _login_last_cleanup
    now = time.time()
    # Periodically clean up old entries to prevent memory leak
    if now - _login_last_cleanup > _LOGIN_WINDOW:
        stale_keys = [k for k, v in _login_attempts.items() if not v or now - v[-1] > _LOGIN_WINDOW]
        for k in stale_keys:
            del _login_attempts[k]
        _login_last_cleanup = now
    _login_attempts[key] = [t for t in _login_attempts[key] if now - t < _LOGIN_WINDOW]
    if len(_login_attempts[key]) >= _LOGIN_MAX_ATTEMPTS:
        raise HTTPException(429, "Слишком много попыток. Подождите 5 минут.")
    _login_attempts[key].append(now)


def sanitize(text: str) -> str:
    """Strip HTML tags to prevent stored XSS."""
    return re.sub(r'<[^>]+>', '', text).strip()


UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

from contextlib import asynccontextmanager
import asyncio

async def _auto_approve_loop():
    """Background task: check overdue approvals every 30 minutes."""
    while True:
        await asyncio.sleep(1800)  # 30 min
        try:
            db = SessionLocal()
            _process_auto_approvals(db)
            db.close()
        except Exception:
            pass

async def _deadline_reminder_loop():
    """Background task: send reminders 1-3 days before deadline."""
    while True:
        await asyncio.sleep(3600)  # every hour
        try:
            db = SessionLocal()
            _process_deadline_reminders(db)
            _process_task_reminders(db)
            db.close()
        except Exception:
            pass

def _process_deadline_reminders(db: Session):
    """Send notifications for documents approaching deadline."""
    now = datetime.now(timezone.utc)
    from datetime import timedelta
    docs = db.query(Document).filter(
        Document.deleted == False,
        Document.status.in_(["draft", "pending"]),
        Document.deadline != "",
    ).all()
    for doc in docs:
        try:
            dl = datetime.strptime(doc.deadline[:10], "%Y-%m-%d").replace(tzinfo=timezone.utc)
        except (ValueError, TypeError):
            continue
        days_left = (dl - now).days
        if days_left < 0 or days_left > 3:
            continue
        # Check if reminder already sent today
        today_str = now.strftime("%Y-%m-%d")
        existing = db.query(Notification).filter(
            Notification.user_id == doc.author_id,
            Notification.doc_id == doc.id,
            Notification.notif_type == "deadline_reminder",
            Notification.created_at >= datetime.strptime(today_str, "%Y-%m-%d").replace(tzinfo=timezone.utc),
        ).first()
        if existing:
            continue
        if days_left == 0:
            msg = f'Сегодня истекает срок: "{doc.title}"'
        elif days_left == 1:
            msg = f'Завтра истекает срок: "{doc.title}"'
        else:
            msg = f'Через {days_left} дня истекает срок: "{doc.title}"'
        db.add(Notification(
            user_id=doc.author_id, notif_type="deadline_reminder",
            title="Напоминание о дедлайне", message=msg, doc_id=doc.id,
        ))
        # Also notify approvers if pending
        if doc.status == "pending":
            for a in doc.approvals:
                if a.status == "pending":
                    db.add(Notification(
                        user_id=a.user_id, notif_type="deadline_reminder",
                        title="Напоминание о дедлайне", message=msg, doc_id=doc.id,
                    ))
    db.commit()


def _process_task_reminders(db: Session):
    """Send notifications for tasks approaching deadline."""
    now = datetime.now(timezone.utc)
    from datetime import timedelta
    tasks = db.query(Task).filter(
        Task.status.in_(["pending", "in_progress"]),
        Task.deadline != "",
    ).all()
    for task in tasks:
        try:
            dl = datetime.strptime(task.deadline[:10], "%Y-%m-%d").replace(tzinfo=timezone.utc)
        except (ValueError, TypeError):
            continue
        days_left = (dl - now).days
        if days_left < 0:
            # Overdue task notification
            today_str = now.strftime("%Y-%m-%d")
            existing = db.query(Notification).filter(
                Notification.user_id == task.assignee_id,
                Notification.notif_type == "task_overdue",
                Notification.doc_id == task.id,
                Notification.created_at >= datetime.strptime(today_str, "%Y-%m-%d").replace(tzinfo=timezone.utc),
            ).first()
            if not existing:
                add_notification(db, task.assignee_id, "task_overdue", "Просроченное поручение",
                                 f'Поручение "{task.title}" просрочено!', task.document_id)
        elif days_left <= 1:
            today_str = now.strftime("%Y-%m-%d")
            existing = db.query(Notification).filter(
                Notification.user_id == task.assignee_id,
                Notification.notif_type == "task_reminder",
                Notification.doc_id == task.id,
                Notification.created_at >= datetime.strptime(today_str, "%Y-%m-%d").replace(tzinfo=timezone.utc),
            ).first()
            if not existing:
                msg = f'Завтра истекает срок поручения: "{task.title}"' if days_left == 1 else f'Сегодня истекает срок поручения: "{task.title}"'
                add_notification(db, task.assignee_id, "task_reminder", "Напоминание", msg, task.document_id)
    db.commit()


# --- Email notifications (SMTP) ---
SMTP_HOST = os.getenv("SMTP_HOST", "")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASS = os.getenv("SMTP_PASS", "")
SMTP_FROM = os.getenv("SMTP_FROM", "")


def send_email(to_addr: str, subject: str, body: str):
    """Send email notification via SMTP (non-blocking)."""
    if not SMTP_HOST or not to_addr:
        return
    import threading
    def _send():
        try:
            import smtplib
            from email.mime.text import MIMEText
            from email.mime.multipart import MIMEMultipart
            msg = MIMEMultipart()
            msg["From"] = SMTP_FROM or SMTP_USER
            msg["To"] = to_addr
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "html", "utf-8"))
            with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
                server.starttls()
                if SMTP_USER and SMTP_PASS:
                    server.login(SMTP_USER, SMTP_PASS)
                server.send_message(msg)
        except Exception:
            pass
    threading.Thread(target=_send, daemon=True).start()


def add_audit(db: Session, user_id: int, user_name: str, action: str,
              entity_type: str = "", entity_id: int = None, details: str = "", ip: str = ""):
    db.add(AuditLog(
        user_id=user_id, user_name=user_name, action=action,
        entity_type=entity_type, entity_id=entity_id, details=details, ip_address=ip,
    ))


@asynccontextmanager
async def lifespan(application):
    run_migrations(engine)
    Base.metadata.create_all(bind=engine)
    _seed_data()
    _renumber_correspondence()
    task = asyncio.create_task(_auto_approve_loop())
    task2 = asyncio.create_task(_deadline_reminder_loop())
    yield
    task.cancel()
    task2.cancel()

app = FastAPI(title="ЭДО API", lifespan=lifespan)



# --- Startup ---
def run_migrations(eng):
    """Add missing columns/tables to existing DB without dropping data."""
    from sqlalchemy import text, inspect
    insp = inspect(eng)
    existing_tables = insp.get_table_names()
    with eng.connect() as conn:
        migrations = [
            ("documents", "extra_fields", "TEXT DEFAULT '{}'"),
            ("documents", "deleted", "BOOLEAN DEFAULT FALSE"),
            ("users", "deputy_id", "INTEGER"),
            ("users", "login", "VARCHAR(20)"),
            ("users", "auto_approve_hours", "INTEGER DEFAULT 0"),
            ("attachments", "filepath", "VARCHAR(1000) DEFAULT ''"),
            ("attachments", "filesize", "INTEGER DEFAULT 0"),
            ("users", "user_status", "VARCHAR(20) DEFAULT 'available'"),
            ("users", "notify_email", "VARCHAR(200) DEFAULT ''"),
            ("users", "notify_telegram", "VARCHAR(100) DEFAULT ''"),
            ("users", "notify_browser", "BOOLEAN DEFAULT TRUE"),
            ("users", "notify_on_approve", "BOOLEAN DEFAULT TRUE"),
            ("users", "notify_on_reject", "BOOLEAN DEFAULT TRUE"),
            ("users", "notify_on_comment", "BOOLEAN DEFAULT TRUE"),
            ("users", "notify_on_task", "BOOLEAN DEFAULT TRUE"),
        ]
        for table, col, col_type in migrations:
            if table in existing_tables:
                cols = [c["name"] for c in insp.get_columns(table)]
                if col not in cols:
                    conn.execute(text(f'ALTER TABLE {table} ADD COLUMN {col} {col_type}'))
        conn.commit()

    # Ensure new tables exist
    for table_name in ["audit_log", "document_templates", "nomenclature_cases"]:
        if table_name not in existing_tables:
            Base.metadata.create_all(bind=eng, tables=[Base.metadata.tables[table_name]])

    # Add case_id column to documents if missing
    if "documents" in existing_tables:
        cols = [c["name"] for c in insp.get_columns("documents")]
        if "case_id" not in cols:
            with eng.connect() as conn:
                conn.execute(text("ALTER TABLE documents ADD COLUMN case_id INTEGER"))
                conn.commit()


def _renumber_correspondence():
    """Перенумеровать входящие и исходящие документы: 1, 2, 3..."""
    db = SessionLocal()
    try:
        for types in [INCOMING_TYPES, OUTGOING_TYPES]:
            docs = db.query(Document).filter(
                Document.doc_type.in_(types)
            ).order_by(Document.created_at).all()
            for i, doc in enumerate(docs, 1):
                new_num = str(i)
                if doc.number != new_num:
                    doc.number = new_num
        db.commit()
    finally:
        db.close()


def _seed_data():
    db = SessionLocal()
    if db.query(Tag).count() == 0:
        for name, color in [("Срочно","red"),("Финансы","green"),("Кадры","blue"),("Продажи","yellow"),("Важно","purple"),("Личное","pink")]:
            db.add(Tag(name=name, color=color))
        db.commit()

    # Seed 3 demo users
    demo_users = [
        {"login": "admedo", "name": "Администратор", "email": "admin@edo.com", "password": "admin123",
         "role": "admin", "department": "Руководство", "position": "Системный администратор", "color": "#2563eb"},
        {"login": "manger", "name": "Менеджер Иванов", "email": "manager@edo.com", "password": "manager123",
         "role": "user", "department": "Управление", "position": "Менеджер проектов", "color": "#16a34a"},
        {"login": "usredo", "name": "Сотрудник Петров", "email": "user@edo.com", "password": "user123",
         "role": "user", "department": "Отдел разработки", "position": "Специалист", "color": "#7c3aed"},
        {"login": "buhgal", "name": "Бухгалтер Смирнова", "email": "buh@edo.com", "password": "buh123",
         "role": "user", "department": "Бухгалтерия", "position": "Главный бухгалтер", "color": "#ea580c"},
    ]
    for u in demo_users:
        existing = db.query(User).filter(User.email == u["email"]).first()
        if not existing:
            db.add(User(
                login=u["login"], name=u["name"], email=u["email"],
                password_hash=hash_password(u["password"]),
                role=u["role"], department=u["department"],
                position=u["position"], color=u["color"],
            ))
        else:
            if not existing.login:
                existing.login = u["login"]
    db.commit()

    # Удалить старых тестовых пользователей без логина
    old_users = db.query(User).filter(User.login == None).all()
    for ou in old_users:
        try:
            db.delete(ou)
            db.commit()
        except Exception:
            db.rollback()
    db.close()


# ============ AUTH ============

@app.post("/api/register", response_model=Token)
def register(data: UserRegister, request: Request, db: Session = Depends(get_db)):
    login_val = data.login.strip().lower()
    if not login_val.isalpha() or len(login_val) != 6:
        raise HTTPException(400, "Логин должен состоять из 6 английских букв")
    if db.query(User).filter(User.login == login_val).first():
        raise HTTPException(400, "Логин уже занят")
    colors = ["#2563eb","#16a34a","#d97706","#7c3aed","#db2777","#059669","#ea580c","#4f46e5"]
    user = User(
        login=login_val, name=sanitize(data.name), email=f"{login_val}@edo.local",
        password_hash=hash_password(data.password),
        department=sanitize(data.department), position=sanitize(data.position),
        color=colors[db.query(User).count() % len(colors)],
    )
    db.add(user)
    db.flush()
    add_audit(db, user.id, user.name, "register", "user", user.id, ip=request.client.host if request.client else "")
    db.commit()
    db.refresh(user)
    token = create_token(user.id)
    return Token(access_token=token, user=UserOut.model_validate(user))


@app.post("/api/login", response_model=Token)
def login(data: UserLogin, request: Request, db: Session = Depends(get_db)):
    login_val = data.login.strip().lower()
    _check_rate_limit(login_val)
    user = db.query(User).filter(User.login == login_val).first()
    if not user or not verify_password(data.password, user.password_hash):
        raise HTTPException(401, "Неверный логин или пароль")
    add_audit(db, user.id, user.name, "login", "user", user.id, ip=request.client.host if request.client else "")
    db.commit()
    token = create_token(user.id)
    return Token(access_token=token, user=UserOut.model_validate(user))


@app.get("/api/me", response_model=UserOut)
def me(user: User = Depends(get_current_user)):
    return UserOut.model_validate(user)


# ============ USERS ============

@app.get("/api/users")
def list_users(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    users = db.query(User).all()
    if user.role == "admin":
        return [UserOut.model_validate(u) for u in users]
    return [UserOutPublic.model_validate(u) for u in users]


@app.post("/api/users", response_model=UserOut)
def create_user(data: UserCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только админ может добавлять сотрудников")
    login_val = data.login.strip().lower()
    if not login_val.isalpha() or len(login_val) != 6:
        raise HTTPException(400, "Логин должен состоять из 6 английских букв")
    if db.query(User).filter(User.login == login_val).first():
        raise HTTPException(400, "Логин уже занят")
    colors = ["#2563eb","#16a34a","#d97706","#7c3aed","#db2777","#059669","#ea580c"]
    new_user = User(
        login=login_val, name=sanitize(data.name), email=f"{login_val}@edo.local",
        password_hash=hash_password(data.password),
        role=data.role, department=sanitize(data.department), position=sanitize(data.position),
        color=colors[db.query(User).count() % len(colors)],
    )
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return UserOut.model_validate(new_user)


@app.put("/api/users/{user_id}/deputy", response_model=UserOut)
def set_deputy(user_id: int, data: DeputySet, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.id != user_id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(404, "Пользователь не найден")
    if data.deputy_id:
        dep = db.query(User).filter(User.id == data.deputy_id).first()
        if not dep:
            raise HTTPException(404, "Заместитель не найден")
        if data.deputy_id == user_id:
            raise HTTPException(400, "Нельзя назначить себя заместителем")
    target.deputy_id = data.deputy_id
    db.commit()
    db.refresh(target)
    return UserOut.model_validate(target)


# ============ PROFILE ============

@app.put("/api/profile")
def update_profile(data: ProfileUpdate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if data.name is not None:
        user.name = sanitize(data.name)
    if data.department is not None:
        user.department = sanitize(data.department)
    if data.position is not None:
        user.position = sanitize(data.position)
    if data.user_status is not None:
        if data.user_status not in ("available", "away", "vacation"):
            raise HTTPException(400, "Неверный статус")
        user.user_status = data.user_status
    if data.notify_email is not None:
        user.notify_email = sanitize(data.notify_email)
    if data.notify_telegram is not None:
        user.notify_telegram = sanitize(data.notify_telegram)
    if data.notify_browser is not None:
        user.notify_browser = data.notify_browser
    if data.notify_on_approve is not None:
        user.notify_on_approve = data.notify_on_approve
    if data.notify_on_reject is not None:
        user.notify_on_reject = data.notify_on_reject
    if data.notify_on_comment is not None:
        user.notify_on_comment = data.notify_on_comment
    if data.notify_on_task is not None:
        user.notify_on_task = data.notify_on_task
    db.commit()
    db.refresh(user)
    return UserOut.model_validate(user)


@app.post("/api/profile/password")
def change_password(data: PasswordChange, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if not verify_password(data.old_password, user.password_hash):
        raise HTTPException(400, "Неверный текущий пароль")
    if len(data.new_password) < 4:
        raise HTTPException(400, "Пароль должен быть не менее 4 символов")
    user.password_hash = hash_password(data.new_password)
    db.commit()
    return {"ok": True}


# ============ TAGS ============

@app.get("/api/tags", response_model=list[TagOut])
def list_tags(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    return [TagOut.model_validate(t) for t in db.query(Tag).all()]


# ============ DOCUMENTS ============

TYPE_PREFIX = {
    "contract":"ДОГ","invoice":"СЧ","order":"ПР","report":"ОТЧ","memo":"СЗ",
    "statement":"ЗАЯ","protocol":"ПРОТ","letter":"ПИС",
    "incoming_letter":"ВХ","outgoing_letter":"ИСХ","nda":"НДА",
    "vacation":"ОТП","trip":"КОМ","purchase":"ЗЗ","job_desc":"ДИ",
    "act":"АКТ","regulation":"ПОЛ",
    "advance_report":"АО","payment_order":"ПП","invoice_tax":"СФ",
    "waybill":"ТН","accounting_memo":"БС","power_of_attorney":"ДОВ",
    "cash_order":"КО",
    "incoming_invoice":"ВХ-СЧ","incoming_act":"ВХ-АКТ","incoming_waybill":"ВХ-ТН",
    "incoming_invoice_tax":"ВХ-СФ","incoming_notification":"ВХ-УВ",
    "incoming_request":"ВХ-ТР","incoming_reconciliation":"ВХ-АС",
    "incoming_contract":"ВХ-ДОГ",
    "outgoing_invoice":"ИСХ-СЧ","outgoing_act":"ИСХ-АКТ","outgoing_waybill":"ИСХ-ТН",
    "outgoing_invoice_tax":"ИСХ-СФ","outgoing_notification":"ИСХ-УВ",
    "outgoing_request":"ИСХ-ТР","outgoing_reconciliation":"ИСХ-АС",
    "outgoing_contract":"ИСХ-ДОГ",
    "other":"ДОК",
}

VALID_DOC_TYPES = set(TYPE_PREFIX.keys())

INCOMING_TYPES = [
    "incoming_letter", "incoming_invoice", "incoming_act", "incoming_waybill",
    "incoming_invoice_tax", "incoming_notification", "incoming_request",
    "incoming_reconciliation", "incoming_contract",
]
OUTGOING_TYPES = [
    "outgoing_letter", "outgoing_invoice", "outgoing_act", "outgoing_waybill",
    "outgoing_invoice_tax", "outgoing_notification", "outgoing_request",
    "outgoing_reconciliation", "outgoing_contract",
]

def gen_number(db: Session, doc_type: str) -> str:
    from sqlalchemy import func
    # Incoming/outgoing documents get simple sequential numbers: 1, 2, 3...
    if doc_type in INCOMING_TYPES:
        count = db.query(func.count(Document.id)).filter(Document.doc_type.in_(INCOMING_TYPES)).scalar() + 1
        number = str(count)
        while db.query(Document).filter(Document.number == number, Document.doc_type.in_(INCOMING_TYPES)).first():
            count += 1
            number = str(count)
        return number
    if doc_type in OUTGOING_TYPES:
        count = db.query(func.count(Document.id)).filter(Document.doc_type.in_(OUTGOING_TYPES)).scalar() + 1
        number = str(count)
        while db.query(Document).filter(Document.number == number, Document.doc_type.in_(OUTGOING_TYPES)).first():
            count += 1
            number = str(count)
        return number
    # Other document types keep prefix-based numbering
    prefix = TYPE_PREFIX.get(doc_type, "ДОК")
    year = datetime.now().year
    count = db.query(func.count(Document.id)).filter(Document.doc_type == doc_type).scalar() + 1
    total = db.query(func.count(Document.id)).scalar() + 1
    number = f"{prefix}-{year}-{str(count).zfill(3)} (№{total})"
    while db.query(Document).filter(Document.number == number).first():
        count += 1
        total += 1
        number = f"{prefix}-{year}-{str(count).zfill(3)} (№{total})"
    return number


def doc_to_out(doc: Document) -> DocumentOut:
    return DocumentOut(
        id=doc.id, number=doc.number or "", title=doc.title,
        description=doc.description or "", content=doc.content or "",
        doc_type=doc.doc_type, status=doc.status, priority=doc.priority or "normal",
        sequential=doc.sequential, deadline=doc.deadline or "",
        extra_fields=json.loads(doc.extra_fields) if isinstance(doc.extra_fields, str) else (doc.extra_fields or {}),
        deleted=doc.deleted or False,
        author_id=doc.author_id, author_name=doc.author_user.name if doc.author_user else "",
        created_at=doc.created_at, updated_at=doc.updated_at,
        approvals=[ApprovalOut(
            id=a.id, user_id=a.user_id, user_name=a.user.name if a.user else "",
            status=a.status, comment=a.comment or "", signature=a.signature or "",
            order_num=a.order_num, decided_at=a.decided_at
        ) for a in sorted(doc.approvals, key=lambda x: x.order_num)],
        comments=[CommentOut(
            id=c.id, user_id=c.user_id, user_name=c.user.name if c.user else "",
            text=c.text, created_at=c.created_at
        ) for c in doc.comments],
        history=[{"id": h.id, "user_name": h.user_name, "text": h.text, "created_at": h.created_at} for h in doc.history],
        versions=[{"id": v.id, "title": v.title, "content": v.content, "user_name": v.user_name, "created_at": v.created_at} for v in doc.versions],
        attachments=[{"id": att.id, "filename": att.filename, "filepath": att.filepath or "", "size": att.size, "filesize": att.filesize or 0} for att in doc.attachments],
        tags=[TagOut.model_validate(t) for t in doc.tags],
        related_doc_ids=[r.id for r in doc.related_docs],
        resolution=ResolutionOut(
            id=doc.resolution.id, user_id=doc.resolution.user_id,
            user_name=doc.resolution.user.name if doc.resolution.user else "",
            text=doc.resolution.text, created_at=doc.resolution.created_at
        ) if doc.resolution else None,
    )


def load_doc(db: Session, doc_id: int) -> Document:
    doc = db.query(Document).options(
        joinedload(Document.author_user),
        joinedload(Document.approvals).joinedload(Approval.user),
        joinedload(Document.comments).joinedload(Comment.user),
        joinedload(Document.history),
        joinedload(Document.versions),
        joinedload(Document.attachments),
        joinedload(Document.tags),
        joinedload(Document.related_docs),
        joinedload(Document.resolution).joinedload(Resolution.user),
    ).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    return doc


def add_history(db: Session, doc: Document, user_name: str, text: str):
    db.add(History(document_id=doc.id, user_name=user_name, text=text))


TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN", "")


def send_telegram(chat_id: str, text: str):
    """Send notification via Telegram Bot API (non-blocking)."""
    if not TELEGRAM_BOT_TOKEN or not chat_id:
        return
    import threading
    def _send():
        try:
            import urllib.request, urllib.parse
            url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
            data = urllib.parse.urlencode({"chat_id": chat_id, "text": text, "parse_mode": "HTML"}).encode()
            urllib.request.urlopen(url, data, timeout=10)
        except Exception:
            pass
    threading.Thread(target=_send, daemon=True).start()


def add_notification(db: Session, user_id: int, notif_type: str, title: str, message: str, doc_id: int = None):
    db.add(Notification(user_id=user_id, notif_type=notif_type, title=title, message=message, doc_id=doc_id))
    try:
        user = db.query(User).filter(User.id == user_id).first()
        if user:
            if user.notify_telegram:
                send_telegram(user.notify_telegram, f"<b>{title}</b>\n{message}")
            if user.notify_email and SMTP_HOST:
                send_email(user.notify_email, f"ЭДО: {title}", f"<h3>{title}</h3><p>{message}</p>")
    except Exception:
        pass


@app.get("/api/documents", response_model=list[DocumentOut])
def list_documents(include_deleted: bool = False, limit: int = 200, offset: int = 0, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    q = db.query(Document).options(
        joinedload(Document.author_user),
        joinedload(Document.approvals).joinedload(Approval.user),
        joinedload(Document.comments).joinedload(Comment.user),
        joinedload(Document.history),
        joinedload(Document.versions),
        joinedload(Document.attachments),
        joinedload(Document.tags),
        joinedload(Document.related_docs),
        joinedload(Document.resolution).joinedload(Resolution.user),
    )
    if not include_deleted:
        q = q.filter(Document.deleted == False)
    if user.role != "admin":
        q = q.filter(
            (Document.author_id == user.id) |
            Document.id.in_(
                db.query(Approval.document_id).filter(Approval.user_id == user.id)
            )
        )
    if limit > 500:
        limit = 500
    docs = q.order_by(Document.updated_at.desc()).offset(offset).limit(limit).all()
    return [doc_to_out(d) for d in docs]


@app.get("/api/documents/trash", response_model=list[DocumentOut])
def list_trash(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    q = db.query(Document).options(
        joinedload(Document.author_user),
        joinedload(Document.approvals).joinedload(Approval.user),
        joinedload(Document.comments).joinedload(Comment.user),
        joinedload(Document.history),
        joinedload(Document.versions),
        joinedload(Document.attachments),
        joinedload(Document.tags),
        joinedload(Document.related_docs),
        joinedload(Document.resolution).joinedload(Resolution.user),
    ).filter(Document.deleted == True)
    if user.role != "admin":
        q = q.filter(Document.author_id == user.id)
    docs = q.order_by(Document.updated_at.desc()).all()
    return [doc_to_out(d) for d in docs]


@app.post("/api/documents", response_model=DocumentOut)
def create_document(data: DocumentCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if data.doc_type not in VALID_DOC_TYPES:
        raise HTTPException(400, f"Неизвестный тип документа: {data.doc_type}")
    number = gen_number(db, data.doc_type)
    doc = Document(
        number=number, title=sanitize(data.title), description=sanitize(data.description),
        content=sanitize(data.content), doc_type=data.doc_type, status=data.status,
        priority=data.priority, sequential=data.sequential,
        deadline=data.deadline, extra_fields=json.dumps(data.extra_fields or {}, ensure_ascii=False),
        author_id=user.id,
    )
    db.add(doc)
    db.flush()

    if data.tag_ids:
        tags = db.query(Tag).filter(Tag.id.in_(data.tag_ids)).all()
        doc.tags = tags

    for att in data.attachments:
        db.add(Attachment(document_id=doc.id, filename=att.get("name",""), size=att.get("size","")))

    if data.related_doc_ids:
        related = db.query(Document).filter(Document.id.in_(data.related_doc_ids)).all()
        doc.related_docs = related

    add_history(db, doc, user.name, "Создан")
    add_audit(db, user.id, user.name, "create", "document", doc.id, f"{doc.doc_type}: {doc.title[:60]}")

    if data.status == "pending" and data.approver_ids:
        valid_uids = [u.id for u in db.query(User.id).filter(User.id.in_(data.approver_ids)).all()]
        if not valid_uids:
            raise HTTPException(400, "Указанные согласователи не найдены")
        for i, uid in enumerate(valid_uids):
            db.add(Approval(document_id=doc.id, user_id=uid, order_num=i))
        add_history(db, doc, user.name, "Отправлен на согласование")
        for uid in valid_uids:
            add_notification(db, uid, "approval_request", "Документ на согласование", f'{user.name}: "{data.title}"', doc.id)

    db.commit()
    return doc_to_out(load_doc(db, doc.id))


def check_doc_access(doc: Document, user: User, db: Session):
    """Check if user has access to document (author, approver, or admin)."""
    if user.role == "admin":
        return
    if doc.author_id == user.id:
        return
    if any(a.user_id == user.id for a in doc.approvals):
        return
    raise HTTPException(403, "Нет доступа к документу")


@app.get("/api/documents/correspondence")
def search_correspondence(
    direction: str = "incoming",
    doc_type: str = "",
    status: str = "",
    q: str = "",
    date_from: str = "",
    date_to: str = "",
    limit: int = 100, offset: int = 0,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    if user.role != "admin" and user.department != "Бухгалтерия":
        raise HTTPException(403, "Доступ только для бухгалтерии и администратора")
    types = INCOMING_TYPES if direction == "incoming" else OUTGOING_TYPES
    query = db.query(Document).options(
        joinedload(Document.author_user),
        joinedload(Document.tags),
    ).filter(Document.deleted == False, Document.doc_type.in_(types))
    if doc_type:
        query = query.filter(Document.doc_type == doc_type)
    if status:
        query = query.filter(Document.status == status)
    if q:
        search = f"%{q}%"
        query = query.filter(
            Document.title.ilike(search) | Document.number.ilike(search) |
            Document.description.ilike(search) | Document.extra_fields.ilike(search)
        )
    if date_from:
        try:
            df = datetime.strptime(date_from, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            query = query.filter(Document.created_at >= df)
        except ValueError:
            pass
    if date_to:
        try:
            dt = datetime.strptime(date_to, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            dt = dt.replace(hour=23, minute=59, second=59)
            query = query.filter(Document.created_at <= dt)
        except ValueError:
            pass
    if limit > 500:
        limit = 500
    docs = query.order_by(Document.created_at.desc()).offset(offset).limit(limit).all()
    results = []
    for d in docs:
        extra = json.loads(d.extra_fields) if isinstance(d.extra_fields, str) else (d.extra_fields or {})
        results.append({
            "id": d.id, "number": d.number or "", "title": d.title,
            "doc_type": d.doc_type, "status": d.status,
            "author_name": d.author_user.name if d.author_user else "",
            "created_at": str(d.created_at), "deadline": d.deadline or "",
            "extra_fields": extra,
        })
    total = query.count() if not offset else len(results)
    return {"results": results, "total": total}


@app.post("/api/documents/bulk")
def bulk_action(data: BulkAction, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if not data.doc_ids:
        raise HTTPException(400, "Не указаны документы")
    if len(data.doc_ids) > 100:
        raise HTTPException(400, "Максимум 100 документов за раз")
    if data.action not in ("delete", "archive", "restore"):
        raise HTTPException(400, "Неизвестное действие")
    docs = db.query(Document).filter(Document.id.in_(data.doc_ids)).all()
    processed = 0
    for doc in docs:
        if doc.author_id != user.id and user.role != "admin":
            continue
        if data.action == "delete":
            doc.deleted = True
            doc.updated_at = datetime.now(timezone.utc)
            add_history(db, doc, user.name, "Перемещён в корзину (массово)")
        elif data.action == "archive":
            doc.status = "archived"
            doc.updated_at = datetime.now(timezone.utc)
            add_history(db, doc, user.name, "В архив (массово)")
        elif data.action == "restore":
            doc.deleted = False
            doc.updated_at = datetime.now(timezone.utc)
            add_history(db, doc, user.name, "Восстановлен (массово)")
        processed += 1
    add_audit(db, user.id, user.name, f"bulk_{data.action}", "document", details=f"{processed} документов")
    db.commit()
    return {"ok": True, "processed": processed}


@app.get("/api/documents/search")
def search_documents(
    q: str = "",
    doc_type: str = "",
    status: str = "",
    author_id: int = 0,
    date_from: str = "",
    date_to: str = "",
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    query = db.query(Document).options(
        joinedload(Document.author_user),
        joinedload(Document.approvals).joinedload(Approval.user),
        joinedload(Document.tags),
        joinedload(Document.resolution).joinedload(Resolution.user),
    ).filter(Document.deleted == False)

    if user.role != "admin":
        query = query.filter(
            (Document.author_id == user.id) |
            Document.id.in_(db.query(Approval.document_id).filter(Approval.user_id == user.id))
        )

    if q:
        search = f"%{q}%"
        # Полнотекстовый поиск: заголовок, содержимое, номер, описание, доп. поля, комментарии
        comment_doc_ids = db.query(Comment.document_id).filter(Comment.text.ilike(search)).subquery()
        query = query.filter(
            Document.title.ilike(search) |
            Document.content.ilike(search) |
            Document.number.ilike(search) |
            Document.description.ilike(search) |
            Document.extra_fields.ilike(search) |
            Document.id.in_(comment_doc_ids)
        )
    if doc_type:
        query = query.filter(Document.doc_type == doc_type)
    if status:
        query = query.filter(Document.status == status)
    if author_id:
        query = query.filter(Document.author_id == author_id)
    if date_from:
        try:
            df = datetime.strptime(date_from, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            query = query.filter(Document.created_at >= df)
        except ValueError:
            pass
    if date_to:
        try:
            dt = datetime.strptime(date_to, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            dt = dt.replace(hour=23, minute=59, second=59)
            query = query.filter(Document.created_at <= dt)
        except ValueError:
            pass

    docs = query.order_by(Document.updated_at.desc()).limit(100).all()

    results = []
    for d in docs:
        results.append({
            "id": d.id,
            "number": d.number or "",
            "title": d.title,
            "doc_type": d.doc_type,
            "status": d.status,
            "author_name": d.author_user.name if d.author_user else "",
            "created_at": str(d.created_at),
            "deadline": d.deadline or "",
        })
    return {"results": results, "total": len(results)}


@app.get("/api/documents/{doc_id}", response_model=DocumentOut)
def get_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    check_doc_access(doc, user, db)
    return doc_to_out(doc)


@app.put("/api/documents/{doc_id}", response_model=DocumentOut)
def update_document(doc_id: int, data: DocumentCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")

    if doc.content != data.content or doc.title != data.title:
        db.add(Version(document_id=doc.id, title=doc.title, content=doc.content, user_name=user.name))

    doc.title = sanitize(data.title)
    doc.description = sanitize(data.description)
    doc.content = sanitize(data.content)
    doc.doc_type = data.doc_type
    doc.status = data.status
    doc.priority = data.priority
    doc.sequential = data.sequential
    doc.deadline = data.deadline
    doc.extra_fields = json.dumps(data.extra_fields or {}, ensure_ascii=False)
    doc.updated_at = datetime.now(timezone.utc)

    tags = db.query(Tag).filter(Tag.id.in_(data.tag_ids)).all() if data.tag_ids else []
    doc.tags = tags

    for att in doc.attachments:
        if att.filepath and os.path.exists(att.filepath):
            os.remove(att.filepath)
        db.delete(att)
    for att in data.attachments:
        db.add(Attachment(document_id=doc.id, filename=att.get("name",""), size=att.get("size","")))

    related = db.query(Document).filter(Document.id.in_(data.related_doc_ids)).all() if data.related_doc_ids else []
    doc.related_docs = related

    add_history(db, doc, user.name, "Отредактирован")

    if data.status == "pending" and data.approver_ids:
        valid_uids = [u.id for u in db.query(User.id).filter(User.id.in_(data.approver_ids)).all()]
        if not valid_uids:
            raise HTTPException(400, "Указанные согласователи не найдены")
        for a in doc.approvals:
            db.delete(a)
        db.flush()
        for i, uid in enumerate(valid_uids):
            db.add(Approval(document_id=doc.id, user_id=uid, order_num=i))
        add_history(db, doc, user.name, "На согласование")
        for uid in valid_uids:
            add_notification(db, uid, "approval_request", "Документ на согласование", f'{user.name}: "{data.title}"', doc.id)

    db.commit()
    return doc_to_out(load_doc(db, doc.id))


# Мягкое удаление (в корзину)
@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    doc.deleted = True
    doc.updated_at = datetime.now(timezone.utc)
    add_history(db, doc, user.name, "Перемещён в корзину")
    add_audit(db, user.id, user.name, "delete", "document", doc.id, doc.title[:60])
    db.commit()
    return {"ok": True}


# Восстановить из корзины
@app.post("/api/documents/{doc_id}/restore")
def restore_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    doc.deleted = False
    doc.updated_at = datetime.now(timezone.utc)
    add_history(db, doc, user.name, "Восстановлен из корзины")
    db.commit()
    return {"ok": True}


# Окончательное удаление
@app.delete("/api/documents/{doc_id}/permanent")
def permanent_delete(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    # Delete uploaded files
    for att in doc.attachments:
        if att.filepath and os.path.exists(att.filepath):
            os.remove(att.filepath)
    db.delete(doc)
    db.commit()
    return {"ok": True}


# ============ APPROVAL ACTIONS ============

def find_approval_for_user(doc, user, db):
    """Find pending approval for user (direct or as deputy)."""
    sorted_approvals = sorted(doc.approvals, key=lambda x: x.order_num)
    # Direct match
    for a in sorted_approvals:
        if a.user_id == user.id and a.status == "pending":
            if doc.sequential:
                for prev in sorted_approvals:
                    if prev.order_num < a.order_num and prev.status != "approved":
                        return None, "Дождитесь предыдущего согласующего"
            return a, None
    # Deputy match
    for a in sorted_approvals:
        if a.status == "pending":
            approver = db.query(User).filter(User.id == a.user_id).first()
            if approver and approver.deputy_id == user.id:
                if doc.sequential:
                    for prev in sorted_approvals:
                        if prev.order_num < a.order_num and prev.status != "approved":
                            return None, "Дождитесь предыдущего согласующего"
                return a, None
    return None, None


@app.post("/api/documents/{doc_id}/approve", response_model=DocumentOut)
def approve_document(doc_id: int, data: ApprovalAction, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.status != "pending":
        raise HTTPException(400, "Документ не на согласовании")

    approval, err = find_approval_for_user(doc, user, db)
    if err:
        raise HTTPException(400, err)
    if not approval:
        raise HTTPException(400, "Вы не можете согласовать этот документ")

    is_deputy = approval.user_id != user.id
    approval.status = "approved"
    approval.comment = data.comment
    approval.signature = secrets.token_hex(32)
    approval.decided_at = datetime.now(timezone.utc)

    if is_deputy:
        approver = db.query(User).filter(User.id == approval.user_id).first()
        add_history(db, doc, user.name, f"ЭЦП (заместитель {approver.name}): {user.name}")
    else:
        add_history(db, doc, user.name, f"ЭЦП: {user.name}")

    db.flush()

    # Notify next approver in sequential mode
    if doc.sequential:
        sorted_approvals = sorted(doc.approvals, key=lambda x: x.order_num)
        for next_a in sorted_approvals:
            if next_a.status == "pending":
                add_notification(db, next_a.user_id, "approval_request", "Ваша очередь", f'Согласуйте "{doc.title}"', doc.id)
                break

    all_approved = all(a.status == "approved" for a in doc.approvals)
    if all_approved:
        doc.status = "approved"
        add_history(db, doc, user.name, "Полностью согласован")
        add_notification(db, doc.author_id, "approved", "Согласован", f'"{doc.title}" согласован', doc.id)

    add_audit(db, user.id, user.name, "approve", "document", doc.id, doc.title[:60])
    doc.updated_at = datetime.now(timezone.utc)
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.post("/api/documents/{doc_id}/reject", response_model=DocumentOut)
def reject_document(doc_id: int, data: ApprovalAction, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.status != "pending":
        raise HTTPException(400, "Документ не на согласовании")

    approval, err = find_approval_for_user(doc, user, db)
    if err:
        raise HTTPException(400, err)
    if not approval:
        raise HTTPException(400, "Вы не можете отклонить этот документ")

    if not data.comment:
        raise HTTPException(400, "Укажите причину отклонения")

    is_deputy = approval.user_id != user.id
    approval.status = "rejected"
    approval.comment = data.comment
    approval.decided_at = datetime.now(timezone.utc)

    doc.status = "rejected"
    doc.updated_at = datetime.now(timezone.utc)

    if is_deputy:
        approver = db.query(User).filter(User.id == approval.user_id).first()
        add_history(db, doc, user.name, f"Отклонён (заместитель {approver.name}): {user.name} — {data.comment}")
    else:
        add_history(db, doc, user.name, f"Отклонён: {user.name} — {data.comment}")
    add_notification(db, doc.author_id, "rejected", "Отклонён", f'{user.name} отклонил "{doc.title}"', doc.id)
    add_audit(db, user.id, user.name, "reject", "document", doc.id, f"{doc.title[:40]}: {data.comment[:40]}")

    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.post("/api/documents/{doc_id}/recall", response_model=DocumentOut)
def recall_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.status != "pending" or doc.author_id != user.id:
        raise HTTPException(400, "Нельзя отозвать")
    doc.status = "draft"
    for a in doc.approvals:
        a.status = "pending"
        a.comment = ""
        a.decided_at = None
        a.signature = ""
    doc.updated_at = datetime.now(timezone.utc)
    add_history(db, doc, user.name, "Отозван с согласования")
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.post("/api/documents/{doc_id}/resend", response_model=DocumentOut)
def resend_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.author_id != user.id:
        raise HTTPException(403, "Нет прав")
    doc.status = "pending"
    for a in doc.approvals:
        a.status = "pending"
        a.comment = ""
        a.decided_at = None
        a.signature = ""
    doc.updated_at = datetime.now(timezone.utc)
    add_history(db, doc, user.name, "Повторно на согласование")
    for a in doc.approvals:
        add_notification(db, a.user_id, "approval_request", "Повторно на согласование", f'{user.name}: "{doc.title}"', doc.id)
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.post("/api/documents/{doc_id}/archive", response_model=DocumentOut)
def archive_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    doc.status = "archived"
    doc.updated_at = datetime.now(timezone.utc)
    add_history(db, doc, user.name, "В архив")
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.post("/api/documents/{doc_id}/copy", response_model=DocumentOut)
def copy_document(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    orig = load_doc(db, doc_id)
    number = gen_number(db, orig.doc_type)
    doc = Document(
        number=number, title=orig.title + " (копия)", description=orig.description,
        content=orig.content, doc_type=orig.doc_type, status="draft",
        priority=orig.priority, deadline="", extra_fields=orig.extra_fields or "{}",
        author_id=user.id,
    )
    db.add(doc)
    db.flush()
    doc.tags = list(orig.tags)
    for att in orig.attachments:
        db.add(Attachment(document_id=doc.id, filename=att.filename, size=att.size))
    add_history(db, doc, user.name, f"Создан (копия из {orig.number})")
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.post("/api/documents/{doc_id}/delegate")
def delegate_approval(doc_id: int, to_user_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    to_user = db.query(User).filter(User.id == to_user_id).first()
    if not to_user:
        raise HTTPException(404, "Пользователь не найден")
    approval, _ = find_approval_for_user(doc, user, db)
    if not approval:
        raise HTTPException(400, "Нет активного согласования")
    approval.user_id = to_user_id
    add_history(db, doc, user.name, f"Делегировано: {user.name} -> {to_user.name}")
    add_notification(db, to_user_id, "approval_request", "Делегировано", f'{user.name} делегировал "{doc.title}"', doc.id)
    doc.updated_at = datetime.now(timezone.utc)
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


# ============ FILE UPLOAD ============

@app.post("/api/documents/{doc_id}/upload")
async def upload_file(doc_id: int, file: UploadFile = File(...), db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    doc_dir = os.path.join(UPLOAD_DIR, str(doc_id))
    os.makedirs(doc_dir, exist_ok=True)
    safe_name = secrets.token_hex(8) + "_" + (file.filename or "file")
    filepath = os.path.join(doc_dir, safe_name)
    content = await file.read()
    max_size = 50 * 1024 * 1024  # 50 MB
    if len(content) > max_size:
        raise HTTPException(400, "Файл слишком большой (макс. 50 МБ)")
    with open(filepath, "wb") as f:
        f.write(content)
    att = Attachment(
        document_id=doc_id, filename=file.filename or "file",
        filepath=filepath, size=str(len(content)),
        filesize=len(content),
    )
    db.add(att)
    add_history(db, doc, user.name, f"Файл загружен: {file.filename}")
    doc.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(att)
    return {"id": att.id, "filename": att.filename, "size": att.size, "filesize": att.filesize}


@app.get("/api/attachments/{att_id}/download")
def download_file(att_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    att = db.query(Attachment).filter(Attachment.id == att_id).first()
    if not att or not att.filepath or not os.path.exists(att.filepath):
        raise HTTPException(404, "Файл не найден")
    doc = load_doc(db, att.document_id)
    check_doc_access(doc, user, db)
    return FileResponse(att.filepath, filename=att.filename)


@app.delete("/api/attachments/{att_id}")
def delete_file(att_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    att = db.query(Attachment).filter(Attachment.id == att_id).first()
    if not att:
        raise HTTPException(404, "Файл не найден")
    doc = db.query(Document).filter(Document.id == att.document_id).first()
    if doc and doc.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав на удаление файла")
    if att.filepath and os.path.exists(att.filepath):
        os.remove(att.filepath)
    db.delete(att)
    db.commit()
    return {"ok": True}


# ============ COMMENTS ============

@app.post("/api/documents/{doc_id}/comments", response_model=DocumentOut)
def add_comment(doc_id: int, data: CommentCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    db.add(Comment(document_id=doc.id, user_id=user.id, text=sanitize(data.text)))
    add_history(db, doc, user.name, "Комментарий: " + sanitize(data.text)[:40])
    if doc.author_id != user.id:
        add_notification(db, doc.author_id, "comment", "Комментарий", f'{user.name}: "{doc.title}"', doc.id)
    doc.updated_at = datetime.now(timezone.utc)
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


# ============ NOTIFICATIONS ============

@app.get("/api/notifications", response_model=list[NotificationOut])
def list_notifications(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    notifs = db.query(Notification).filter(Notification.user_id == user.id).order_by(Notification.created_at.desc()).limit(200).all()
    return [NotificationOut.model_validate(n) for n in notifs]


@app.post("/api/notifications/read-all")
def read_all_notifications(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    db.query(Notification).filter(Notification.user_id == user.id, Notification.read == False).update({"read": True})
    db.commit()
    return {"ok": True}


@app.post("/api/notifications/{notif_id}/read")
def read_notification(notif_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    n = db.query(Notification).filter(Notification.id == notif_id, Notification.user_id == user.id).first()
    if n:
        n.read = True
        db.commit()
    return {"ok": True}


# ============ TASKS (Поручения) ============

@app.get("/api/tasks", response_model=list[TaskOut])
def list_tasks(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    tasks = db.query(Task).filter(
        (Task.author_id == user.id) | (Task.assignee_id == user.id)
    ).order_by(Task.created_at.desc()).all()
    result = []
    for t in tasks:
        out = TaskOut.model_validate(t)
        out.author_name = t.author.name if t.author else ""
        out.assignee_name = t.assignee.name if t.assignee else ""
        result.append(out)
    return result


@app.post("/api/tasks", response_model=TaskOut)
def create_task(data: TaskCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    assignee = db.query(User).filter(User.id == data.assignee_id).first()
    if not assignee:
        raise HTTPException(404, "Исполнитель не найден")
    task = Task(
        title=sanitize(data.title), description=sanitize(data.description),
        document_id=data.document_id, author_id=user.id,
        assignee_id=data.assignee_id, priority=data.priority,
        deadline=data.deadline,
    )
    db.add(task)
    db.flush()
    add_notification(db, data.assignee_id, "task", "Новое поручение",
                     f'{user.name}: "{data.title}"', data.document_id)
    db.commit()
    db.refresh(task)
    out = TaskOut.model_validate(task)
    out.author_name = user.name
    out.assignee_name = assignee.name
    return out


@app.put("/api/tasks/{task_id}", response_model=TaskOut)
def update_task(task_id: int, data: TaskUpdate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    task = db.query(Task).filter(Task.id == task_id).first()
    if not task:
        raise HTTPException(404, "Поручение не найдено")
    if task.author_id != user.id and task.assignee_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    if data.title is not None:
        task.title = sanitize(data.title)
    if data.description is not None:
        task.description = sanitize(data.description)
    if data.status is not None:
        old_status = task.status
        task.status = data.status
        if data.status == "completed" and old_status != "completed":
            add_notification(db, task.author_id, "task_done", "Поручение выполнено",
                             f'"{task.title}" выполнено', task.document_id)
    if data.priority is not None:
        task.priority = data.priority
    if data.deadline is not None:
        task.deadline = data.deadline
    task.updated_at = datetime.now(timezone.utc)
    db.commit()
    db.refresh(task)
    out = TaskOut.model_validate(task)
    out.author_name = task.author.name if task.author else ""
    out.assignee_name = task.assignee.name if task.assignee else ""
    return out


@app.delete("/api/tasks/{task_id}")
def delete_task(task_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    task = db.query(Task).filter(Task.id == task_id).first()
    if not task:
        raise HTTPException(404, "Поручение не найдено")
    if task.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    db.delete(task)
    db.commit()
    return {"ok": True}


# ============ DASHBOARD ============

@app.get("/api/dashboard")
def get_dashboard(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    from sqlalchemy import func

    # All docs visible to user
    q = db.query(Document).filter(Document.deleted == False)
    if user.role != "admin":
        q = q.filter(
            (Document.author_id == user.id) |
            Document.id.in_(db.query(Approval.document_id).filter(Approval.user_id == user.id))
        )
    docs = q.all()

    total = len(docs)
    by_status = {}
    for d in docs:
        by_status[d.status] = by_status.get(d.status, 0) + 1
    by_type = {}
    for d in docs:
        by_type[d.doc_type] = by_type.get(d.doc_type, 0) + 1

    # Pending approvals for this user
    pending_approvals = []
    for d in docs:
        if d.status == "pending":
            for a in d.approvals:
                if a.user_id == user.id and a.status == "pending":
                    pending_approvals.append({"doc_id": d.id, "title": d.title, "author": d.author_user.name if d.author_user else "", "created_at": str(d.created_at)})
                    break

    # Overdue docs
    now = datetime.now(timezone.utc)
    overdue = []
    for d in docs:
        if d.deadline and d.status not in ("archived", "approved", "resolved"):
            try:
                dl = datetime.fromisoformat(d.deadline.replace("Z", "+00:00")) if "T" in d.deadline else datetime.strptime(d.deadline[:10], "%Y-%m-%d").replace(tzinfo=timezone.utc)
                if dl < now:
                    overdue.append({"doc_id": d.id, "title": d.title, "deadline": d.deadline})
            except (ValueError, TypeError):
                pass

    # Recent docs
    recent = sorted(docs, key=lambda x: x.created_at or datetime.min.replace(tzinfo=timezone.utc), reverse=True)[:5]
    recent_out = [{"id": d.id, "number": d.number, "title": d.title, "status": d.status, "created_at": str(d.created_at)} for d in recent]

    # My tasks
    my_tasks = db.query(Task).filter(
        Task.assignee_id == user.id, Task.status.in_(["pending", "in_progress"])
    ).order_by(Task.created_at.desc()).limit(5).all()
    tasks_out = [{"id": t.id, "title": t.title, "status": t.status, "deadline": t.deadline, "priority": t.priority} for t in my_tasks]

    # Overdue tasks
    overdue_tasks = []
    all_my_tasks = db.query(Task).filter(
        Task.assignee_id == user.id, Task.status.in_(["pending", "in_progress"]),
        Task.deadline != "",
    ).all()
    for t in all_my_tasks:
        try:
            dl = datetime.fromisoformat(t.deadline.replace("Z", "+00:00")) if "T" in t.deadline else datetime.strptime(t.deadline[:10], "%Y-%m-%d").replace(tzinfo=timezone.utc)
            if dl < now:
                overdue_tasks.append({"id": t.id, "title": t.title, "deadline": t.deadline, "priority": t.priority})
        except (ValueError, TypeError):
            pass

    return {
        "total": total,
        "by_status": by_status,
        "by_type": by_type,
        "pending_approvals": pending_approvals,
        "overdue": overdue,
        "overdue_tasks": overdue_tasks,
        "recent": recent_out,
        "my_tasks": tasks_out,
    }


# ============ RESOLUTION ============

@app.post("/api/documents/{doc_id}/resolution", response_model=DocumentOut)
def create_resolution(doc_id: int, data: ResolutionCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    doc = load_doc(db, doc_id)
    if doc.status != "approved":
        raise HTTPException(400, "Резолюция возможна только для согласованных документов")

    # Only last approver or admin can create resolution
    sorted_approvals = sorted(doc.approvals, key=lambda x: x.order_num)
    last_approver_id = sorted_approvals[-1].user_id if sorted_approvals else None
    if user.id != last_approver_id and user.role != "admin":
        raise HTTPException(403, "Резолюцию может создать только последний согласующий или администратор")

    if doc.resolution:
        raise HTTPException(400, "Резолюция уже создана")

    if not data.text.strip():
        raise HTTPException(400, "Текст резолюции не может быть пустым")

    resolution = Resolution(document_id=doc.id, user_id=user.id, text=sanitize(data.text))
    db.add(resolution)
    doc.status = "resolved"
    doc.updated_at = datetime.now(timezone.utc)
    add_history(db, doc, user.name, f"Резолюция: {data.text[:60]}")
    add_notification(db, doc.author_id, "resolution", "Резолюция", f'{user.name} вынес резолюцию по "{doc.title}"', doc.id)
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


# ============ AUTO-APPROVE & DEPUTY ESCALATION ============

@app.post("/api/auto-approve/run")
def run_auto_approve(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    """Manually trigger auto-approval check. Can also be called by cron/scheduler."""
    if user.role != "admin":
        raise HTTPException(403, "Только администратор")
    count = _process_auto_approvals(db)
    return {"processed": count}


def _process_auto_approvals(db: Session) -> int:
    """Check overdue approvals: auto-approve or escalate to deputy."""
    now = datetime.now(timezone.utc)
    pending_docs = db.query(Document).filter(
        Document.status == "pending", Document.deleted == False
    ).options(
        joinedload(Document.approvals).joinedload(Approval.user),
        joinedload(Document.author_user),
    ).all()

    count = 0

    # Auto-delegate approvals when user is on vacation
    for doc in pending_docs:
        for approval in doc.approvals:
            if approval.status != "pending":
                continue
            approver = approval.user
            if not approver:
                continue
            if approver.user_status == "vacation" and approver.deputy_id:
                deputy = db.query(User).filter(User.id == approver.deputy_id).first()
                if deputy:
                    existing_delegation = db.query(Approval).filter(
                        Approval.document_id == doc.id, Approval.user_id == deputy.id
                    ).first()
                    if not existing_delegation:
                        approval.user_id = deputy.id
                        add_history(db, doc, "Система", f"Автоделегирование: {approver.name} (отпуск) → {deputy.name}")
                        add_notification(db, deputy.id, "delegation", "Делегирование",
                                         f'Документ "{doc.title}" делегирован вам (заместитель {approver.name})', doc.id)
                        count += 1

    for doc in pending_docs:
        if not doc.deadline:
            continue
        try:
            dl = datetime.strptime(doc.deadline[:10], "%Y-%m-%d").replace(tzinfo=timezone.utc)
        except (ValueError, TypeError):
            continue
        if dl >= now:
            continue

        # Document is overdue - check each pending approval
        for approval in sorted(doc.approvals, key=lambda a: a.order_num):
            if approval.status != "pending":
                continue

            approver = db.query(User).filter(User.id == approval.user_id).first()
            if not approver:
                continue

            # Auto-approve if user has auto_approve_hours set and time exceeded
            if approver.auto_approve_hours and approver.auto_approve_hours > 0:
                hours_overdue = (now - dl).total_seconds() / 3600
                if hours_overdue >= approver.auto_approve_hours:
                    approval.status = "approved"
                    approval.comment = "Автосогласование (превышен таймаут)"
                    approval.signature = secrets.token_hex(32)
                    approval.decided_at = now
                    add_history(db, doc, "Система", f"Автосогласование: {approver.name} (таймаут {approver.auto_approve_hours}ч)")
                    count += 1

                    # Check if all approved
                    all_approved = all(a.status == "approved" for a in doc.approvals)
                    if all_approved:
                        doc.status = "approved"
                        doc.updated_at = now
                        add_history(db, doc, "Система", "Полностью согласован (авто)")
                        add_notification(db, doc.author_id, "approved", "Согласован (авто)", f'"{doc.title}" согласован автоматически', doc.id)
                    continue

            # Escalate to deputy if set
            if approver.deputy_id:
                deputy = db.query(User).filter(User.id == approver.deputy_id).first()
                if deputy:
                    # Notify deputy about overdue approval
                    existing = db.query(Notification).filter(
                        Notification.user_id == deputy.id,
                        Notification.doc_id == doc.id,
                        Notification.notif_type == "deputy_escalation",
                    ).first()
                    if not existing:
                        add_notification(db, deputy.id, "deputy_escalation",
                                         "Эскалация: требуется согласование",
                                         f'Документ "{doc.title}" просрочен. Вы заместитель {approver.name}.', doc.id)
                        add_history(db, doc, "Система", f"Эскалация заместителю: {deputy.name} (за {approver.name})")
                        count += 1

            # If sequential, only process first pending
            if doc.sequential:
                break

    db.commit()
    return count


@app.put("/api/users/{user_id}/auto-approve")
def set_auto_approve(user_id: int, hours: int = 0, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.id != user_id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    target = db.query(User).filter(User.id == user_id).first()
    if not target:
        raise HTTPException(404, "Пользователь не найден")
    if hours < 0 or hours > 720:
        raise HTTPException(400, "Таймаут должен быть от 0 до 720 часов")
    target.auto_approve_hours = hours
    db.commit()
    return {"ok": True, "auto_approve_hours": hours}


# ============ ROUTES ============

@app.get("/api/routes", response_model=list[RouteOut])
def list_routes(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    return [RouteOut.model_validate(r) for r in db.query(ApprovalRoute).all()]


@app.post("/api/routes", response_model=RouteOut)
def create_route(data: RouteCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только админ может создавать маршруты")
    route = ApprovalRoute(name=sanitize(data.name), user_ids=",".join(str(x) for x in data.user_ids), sequential=data.sequential)
    db.add(route)
    db.commit()
    db.refresh(route)
    return RouteOut.model_validate(route)


@app.delete("/api/routes/{route_id}")
def delete_route(route_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только админ может удалять маршруты")
    route = db.query(ApprovalRoute).filter(ApprovalRoute.id == route_id).first()
    if route:
        db.delete(route)
        db.commit()
    return {"ok": True}


# ============ AUDIT LOG ============

@app.get("/api/audit-log", response_model=list[AuditLogOut])
def list_audit_log(
    limit: int = 100, offset: int = 0,
    action: str = "", user_id: int = 0,
    db: Session = Depends(get_db), user: User = Depends(get_current_user),
):
    if user.role != "admin":
        raise HTTPException(403, "Только администратор")
    q = db.query(AuditLog)
    if action:
        q = q.filter(AuditLog.action == action)
    if user_id:
        q = q.filter(AuditLog.user_id == user_id)
    if limit > 500:
        limit = 500
    logs = q.order_by(AuditLog.created_at.desc()).offset(offset).limit(limit).all()
    return [AuditLogOut.model_validate(l) for l in logs]


# ============ DOCUMENT TEMPLATES ============

@app.get("/api/templates")
def list_templates(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    try:
        templates = db.query(DocumentTemplate).filter(
            (DocumentTemplate.is_public == True) | (DocumentTemplate.author_id == user.id)
        ).order_by(DocumentTemplate.created_at.desc()).all()
    except Exception:
        Base.metadata.create_all(bind=engine)
        db.rollback()
        templates = db.query(DocumentTemplate).filter(
            (DocumentTemplate.is_public == True) | (DocumentTemplate.author_id == user.id)
        ).order_by(DocumentTemplate.created_at.desc()).all()
    result = []
    for t in templates:
        extra = {}
        try:
            extra = json.loads(t.extra_fields_template) if isinstance(t.extra_fields_template, str) else (t.extra_fields_template or {})
        except Exception:
            pass
        result.append({
            "id": t.id,
            "name": t.name,
            "doc_type": t.doc_type,
            "title_template": t.title_template or "",
            "description_template": t.description_template or "",
            "content_template": t.content_template or "",
            "extra_fields_template": extra,
            "priority": t.priority or "normal",
            "approver_ids": t.approver_ids or "",
            "sequential": t.sequential,
            "author_id": t.author_id,
            "author_name": t.author.name if t.author else "",
            "is_public": t.is_public,
            "created_at": t.created_at.isoformat() if t.created_at else "",
        })
    return result


@app.post("/api/templates")
def create_template(data: TemplateCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    tmpl = DocumentTemplate(
        name=sanitize(data.name),
        doc_type=data.doc_type,
        title_template=sanitize(data.title_template),
        description_template=sanitize(data.description_template),
        content_template=sanitize(data.content_template),
        extra_fields_template=json.dumps(data.extra_fields_template or {}, ensure_ascii=False),
        priority=data.priority,
        approver_ids=",".join(str(x) for x in data.approver_ids),
        sequential=data.sequential,
        author_id=user.id,
        is_public=data.is_public,
    )
    db.add(tmpl)
    add_audit(db, user.id, user.name, "create", "template", details=data.name[:60])
    db.commit()
    db.refresh(tmpl)
    extra = {}
    try:
        extra = json.loads(tmpl.extra_fields_template) if isinstance(tmpl.extra_fields_template, str) else {}
    except Exception:
        pass
    return {
        "id": tmpl.id, "name": tmpl.name, "doc_type": tmpl.doc_type,
        "title_template": tmpl.title_template or "", "description_template": tmpl.description_template or "",
        "content_template": tmpl.content_template or "", "extra_fields_template": extra,
        "priority": tmpl.priority or "normal", "approver_ids": tmpl.approver_ids or "",
        "sequential": tmpl.sequential, "author_id": tmpl.author_id,
        "author_name": user.name, "is_public": tmpl.is_public,
        "created_at": tmpl.created_at.isoformat() if tmpl.created_at else "",
    }


@app.post("/api/templates/{tmpl_id}/apply", response_model=DocumentOut)
def apply_template(tmpl_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    tmpl = db.query(DocumentTemplate).filter(DocumentTemplate.id == tmpl_id).first()
    if not tmpl:
        raise HTTPException(404, "Шаблон не найден")
    if not tmpl.is_public and tmpl.author_id != user.id:
        raise HTTPException(403, "Нет доступа к шаблону")

    number = gen_number(db, tmpl.doc_type)
    doc = Document(
        number=number,
        title=tmpl.title_template or tmpl.name,
        description=tmpl.description_template or "",
        content=tmpl.content_template or "",
        doc_type=tmpl.doc_type,
        status="draft",
        priority=tmpl.priority,
        sequential=tmpl.sequential,
        deadline="",
        extra_fields=tmpl.extra_fields_template or "{}",
        author_id=user.id,
    )
    db.add(doc)
    db.flush()
    add_history(db, doc, user.name, f"Создан из шаблона: {tmpl.name}")
    add_audit(db, user.id, user.name, "apply_template", "document", doc.id, tmpl.name[:60])

    # Add approvers from template
    if tmpl.approver_ids:
        ids = [int(x) for x in tmpl.approver_ids.split(",") if x.strip()]
        for i, uid in enumerate(ids):
            db.add(Approval(document_id=doc.id, user_id=uid, order_num=i))

    db.commit()
    return doc_to_out(load_doc(db, doc.id))


@app.delete("/api/templates/{tmpl_id}")
def delete_template(tmpl_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    tmpl = db.query(DocumentTemplate).filter(DocumentTemplate.id == tmpl_id).first()
    if not tmpl:
        raise HTTPException(404, "Шаблон не найден")
    if tmpl.author_id != user.id and user.role != "admin":
        raise HTTPException(403, "Нет прав")
    db.delete(tmpl)
    add_audit(db, user.id, user.name, "delete", "template", tmpl_id, tmpl.name[:60])
    db.commit()
    return {"ok": True}


# ============ DEADLINE REMINDERS (manual trigger) ============

@app.post("/api/reminders/check")
def check_reminders(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только администратор")
    _process_deadline_reminders(db)
    return {"ok": True}


# ============ STATIC ============

# ============ EXPORT (PDF / DOCX) ============

DOC_TYPE_LABELS = {
    "contract":"Договор","invoice":"Счёт на оплату","order":"Приказ","report":"Отчёт",
    "memo":"Служебная записка","statement":"Заявление","protocol":"Протокол","letter":"Письмо",
    "incoming_letter":"Входящее письмо","outgoing_letter":"Исходящее письмо",
    "vacation":"Заявление на отпуск","trip":"Командировка","purchase":"Заявка на закупку",
    "job_desc":"Должностная инструкция","act":"Акт выполненных работ","regulation":"Положение",
    "nda":"NDA","advance_report":"Авансовый отчёт","payment_order":"Платёжное поручение",
    "invoice_tax":"Счёт-фактура","waybill":"Товарная накладная",
    "accounting_memo":"Бухгалтерская справка","power_of_attorney":"Доверенность",
    "cash_order":"Кассовый ордер",
    "incoming_invoice":"Вх. счёт на оплату","incoming_act":"Вх. акт вып. работ",
    "incoming_waybill":"Вх. товарная накладная","incoming_invoice_tax":"Вх. счёт-фактура",
    "incoming_notification":"Вх. уведомление","incoming_request":"Вх. требование",
    "incoming_reconciliation":"Вх. акт сверки","incoming_contract":"Вх. договор",
    "outgoing_invoice":"Исх. счёт на оплату","outgoing_act":"Исх. акт вып. работ",
    "outgoing_waybill":"Исх. товарная накладная","outgoing_invoice_tax":"Исх. счёт-фактура",
    "outgoing_notification":"Исх. уведомление","outgoing_request":"Исх. требование",
    "outgoing_reconciliation":"Исх. акт сверки","outgoing_contract":"Исх. договор",
    "other":"Прочее",
}
STATUS_LABELS = {"draft":"Черновик","pending":"На согласовании","approved":"Согласован","rejected":"Отклонён","resolved":"Исполнен","archived":"Архив"}
PRIORITY_LABELS = {"low":"Низкий","normal":"Обычный","high":"Высокий","urgent":"Срочный"}


@app.get("/api/documents/{doc_id}/export/docx")
def export_docx(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import tempfile

    doc = load_doc(db, doc_id)
    check_doc_access(doc, user, db)
    d = doc_to_out(doc)

    dx = DocxDocument()
    style = dx.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header
    h = dx.add_heading('', level=1)
    run = h.add_run(d.title)
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 0, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    dx.add_paragraph('')
    meta = [
        ('Номер', d.number), ('Тип', DOC_TYPE_LABELS.get(d.doc_type, d.doc_type)),
        ('Статус', STATUS_LABELS.get(d.status, d.status)),
        ('Приоритет', PRIORITY_LABELS.get(d.priority, d.priority)),
        ('Автор', d.author_name), ('Дедлайн', d.deadline or '—'),
        ('Создан', d.created_at.strftime('%d.%m.%Y %H:%M') if d.created_at else ''),
    ]
    for label, value in meta:
        p = dx.add_paragraph()
        p.add_run(f'{label}: ').bold = True
        p.add_run(str(value))

    if d.description:
        dx.add_paragraph('')
        p = dx.add_paragraph()
        p.add_run('Описание: ').bold = True
        p.add_run(d.description)

    # Extra fields
    if d.extra_fields:
        dx.add_paragraph('')
        p = dx.add_paragraph()
        p.add_run('Дополнительные поля:').bold = True
        for k, v in d.extra_fields.items():
            if v:
                p = dx.add_paragraph()
                p.add_run(f'  {k}: ').bold = True
                p.add_run(str(v))

    # Content
    dx.add_paragraph('')
    h2 = dx.add_heading('', level=2)
    h2.add_run('Содержание').font.size = Pt(14)
    for line in (d.content or '').split('\n'):
        dx.add_paragraph(line)

    # Approvals
    if d.approvals:
        dx.add_paragraph('')
        h3 = dx.add_heading('', level=2)
        h3.add_run('Согласование').font.size = Pt(14)
        for a in d.approvals:
            status_text = STATUS_LABELS.get(a.status, a.status)
            p = dx.add_paragraph()
            p.add_run(f'{a.user_name}: ').bold = True
            p.add_run(f'{status_text}')
            if a.comment:
                p.add_run(f' — {a.comment}')
            if a.decided_at:
                p.add_run(f' ({a.decided_at.strftime("%d.%m.%Y %H:%M")})')

    # Comments
    if d.comments:
        dx.add_paragraph('')
        h3 = dx.add_heading('', level=2)
        h3.add_run('Комментарии').font.size = Pt(14)
        for c in d.comments:
            p = dx.add_paragraph()
            p.add_run(f'{c.user_name}: ').bold = True
            p.add_run(c.text)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    dx.save(tmp.name)
    tmp.close()
    safe_title = "".join(c for c in d.title[:40] if c.isalnum() or c in ' _-').strip() or 'document'
    from starlette.background import BackgroundTask
    return FileResponse(tmp.name, filename=f'{d.number} {safe_title}.docx',
                        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        background=BackgroundTask(os.unlink, tmp.name))


@app.get("/api/documents/{doc_id}/export/pdf")
def export_pdf(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    from fpdf import FPDF
    from fastapi.responses import JSONResponse
    import tempfile, traceback

    def clean(text):
        if not text:
            return ""
        out = []
        for ch in str(text):
            cp = ord(ch)
            if 0xD800 <= cp <= 0xDFFF:
                continue
            elif cp == 0xA0:
                out.append(' ')
            elif cp < 0x20 and cp not in (0x0A, 0x0D, 0x09):
                continue
            else:
                out.append(ch)
        return ''.join(out)

    doc = load_doc(db, doc_id)
    check_doc_access(doc, user, db)
    d = doc_to_out(doc)

    try:
        # Pre-clean all string fields
        d.title = clean(d.title)
        d.number = clean(d.number)
        d.description = clean(d.description)
        d.content = clean(d.content)
        d.author_name = clean(d.author_name)
        d.deadline = clean(d.deadline)
        for a in d.approvals:
            a.user_name = clean(a.user_name)
            a.comment = clean(a.comment)
        for c in d.comments:
            c.user_name = clean(c.user_name)
            c.text = clean(c.text)
        if d.extra_fields:
            d.extra_fields = {clean(str(k)): clean(str(v)) for k, v in d.extra_fields.items()}

        # Find DejaVu font
        font_path = font_bold_path = None
        for sp in ["/usr/share/fonts/truetype/dejavu", "/usr/share/fonts/dejavu", "/usr/share/fonts/TTF"]:
            fp = os.path.join(sp, "DejaVuSans.ttf")
            fb = os.path.join(sp, "DejaVuSans-Bold.ttf")
            if os.path.exists(fp):
                font_path, font_bold_path = fp, fb
                break
        if not font_path:
            font_dir = os.path.join(os.path.dirname(__file__), "fonts")
            os.makedirs(font_dir, exist_ok=True)
            fp = os.path.join(font_dir, "DejaVuSans.ttf")
            fb = os.path.join(font_dir, "DejaVuSans-Bold.ttf")
            if not os.path.exists(fp):
                try:
                    import urllib.request
                    base = "https://raw.githubusercontent.com/dejavu-fonts/dejavu-fonts/master/ttf/"
                    urllib.request.urlretrieve(base + "DejaVuSans.ttf", fp)
                    urllib.request.urlretrieve(base + "DejaVuSans-Bold.ttf", fb)
                except Exception:
                    pass
            if os.path.exists(fp):
                font_path, font_bold_path = fp, fb

        pdf = FPDF()
        pdf.add_page()
        if font_path and os.path.exists(font_path):
            pdf.add_font('DejaVu', '', font_path)
            pdf.add_font('DejaVu', 'B', font_bold_path if font_bold_path and os.path.exists(font_bold_path) else font_path)
            font_name = 'DejaVu'
        else:
            font_name = 'Helvetica'
        pdf.set_auto_page_break(auto=True, margin=15)

        # Title
        pdf.set_font(font_name, 'B', 16)
        pdf.multi_cell(0, 10, d.title, align='C')
        pdf.ln(5)

        # Meta
        def add_field(label, value):
            pdf.set_font(font_name, 'B', 10)
            pdf.cell(40, 7, f'{label}:', ln=0)
            pdf.set_font(font_name, '', 10)
            pdf.cell(0, 7, clean(str(value)), ln=1)

        add_field('Номер', d.number)
        add_field('Тип', DOC_TYPE_LABELS.get(d.doc_type, d.doc_type))
        add_field('Статус', STATUS_LABELS.get(d.status, d.status))
        add_field('Приоритет', PRIORITY_LABELS.get(d.priority, d.priority))
        add_field('Автор', d.author_name)
        add_field('Дедлайн', d.deadline or '—')
        add_field('Создан', d.created_at.strftime('%d.%m.%Y %H:%M') if d.created_at else '')

        if d.description:
            pdf.ln(3)
            add_field('Описание', d.description)

        # Extra fields
        if d.extra_fields:
            pdf.ln(5)
            pdf.set_font(font_name, 'B', 12)
            pdf.cell(0, 8, 'Дополнительные поля', ln=1)
            for k, v in d.extra_fields.items():
                if v:
                    add_field(k, v)

        # Content
        pdf.ln(5)
        pdf.set_font(font_name, 'B', 12)
        pdf.cell(0, 8, 'Содержание', ln=1)
        pdf.set_font(font_name, '', 10)
        pdf.set_x(pdf.l_margin)
        pw = pdf.w - pdf.l_margin - pdf.r_margin
        pdf.multi_cell(pw, 6, d.content or '')

        # Approvals
        if d.approvals:
            pdf.ln(5)
            pdf.set_font(font_name, 'B', 12)
            pdf.cell(0, 8, 'Согласование', ln=1)
            for a in d.approvals:
                pdf.set_font(font_name, '', 10)
                line = a.user_name + ': ' + STATUS_LABELS.get(a.status, a.status)
                if a.comment:
                    line += f' — {a.comment}'
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(pw, 7, line)

        # Comments
        if d.comments:
            pdf.ln(5)
            pdf.set_font(font_name, 'B', 12)
            pdf.cell(0, 8, 'Комментарии', ln=1)
            for c in d.comments:
                pdf.set_font(font_name, 'B', 10)
                pdf.cell(0, 7, c.user_name + ':', ln=1)
                pdf.set_font(font_name, '', 10)
                pdf.set_x(pdf.l_margin)
                pdf.multi_cell(pw, 6, c.text)
                pdf.ln(2)

        # Organization stamp (circular) on approved documents
        if d.status in ("approved", "resolved") and ORG_NAME:
            try:
                import math
                stamp_x = pdf.w - 60
                stamp_y = pdf.get_y() + 10
                if stamp_y + 40 > pdf.h - 15:
                    pdf.add_page()
                    stamp_y = 30
                r = 18
                cx, cy = stamp_x + r, stamp_y + r
                # Draw outer circle
                pdf.set_draw_color(0, 100, 200)
                pdf.set_line_width(0.8)
                for i in range(360):
                    a1 = math.radians(i)
                    a2 = math.radians(i + 1)
                    pdf.line(cx + r * math.cos(a1), cy + r * math.sin(a1),
                             cx + r * math.cos(a2), cy + r * math.sin(a2))
                # Inner circle
                r2 = r - 3
                for i in range(360):
                    a1 = math.radians(i)
                    a2 = math.radians(i + 1)
                    pdf.line(cx + r2 * math.cos(a1), cy + r2 * math.sin(a1),
                             cx + r2 * math.cos(a2), cy + r2 * math.sin(a2))
                # Center text
                pdf.set_font(font_name, 'B', 6)
                pdf.set_text_color(0, 100, 200)
                short_name = ORG_NAME[:20]
                tw = pdf.get_string_width(short_name)
                pdf.text(cx - tw / 2, cy - 2, short_name)
                pdf.set_font(font_name, '', 5)
                if ORG_INN:
                    inn_text = f"ИНН {ORG_INN}"
                    tw2 = pdf.get_string_width(inn_text)
                    pdf.text(cx - tw2 / 2, cy + 3, inn_text)
                date_text = d.created_at.strftime('%d.%m.%Y') if d.created_at else ""
                if date_text:
                    tw3 = pdf.get_string_width(date_text)
                    pdf.text(cx - tw3 / 2, cy + 7, date_text)
                pdf.set_text_color(0, 0, 0)
                pdf.set_draw_color(0, 0, 0)
            except Exception:
                pass

        # Watermark on draft/pending documents
        if d.status in ("draft", "pending"):
            try:
                page_count = pdf.page
                for pg in range(1, page_count + 1):
                    pdf.page = pg
                    pdf.set_font(font_name, 'B', 40)
                    pdf.set_text_color(220, 220, 220)
                    wm_text = "ЧЕРНОВИК" if d.status == "draft" else "НА СОГЛАСОВАНИИ"
                    tw = pdf.get_string_width(wm_text)
                    pdf.text((pdf.w - tw) / 2, pdf.h / 2, wm_text)
                pdf.set_text_color(0, 0, 0)
            except Exception:
                pass

        # QR code in PDF
        try:
            import qrcode as qr_lib
            import io as qr_io
            host = os.getenv("RENDER_EXTERNAL_URL", os.getenv("BASE_URL", ""))
            qr_url = f"{host}/#doc/{doc_id}" if host else f"DOC-{d.number}"
            qr = qr_lib.QRCode(version=1, box_size=4, border=1)
            qr.add_data(qr_url)
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white")
            qr_buf = qr_io.BytesIO()
            qr_img.save(qr_buf, format="PNG")
            qr_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            qr_tmp.write(qr_buf.getvalue())
            qr_tmp.close()
            if pdf.get_y() + 30 > pdf.h - 15:
                pdf.add_page()
            pdf.ln(5)
            pdf.image(qr_tmp.name, x=pdf.l_margin, y=pdf.get_y(), w=25, h=25)
            os.unlink(qr_tmp.name)
        except Exception:
            pass

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        pdf.output(tmp.name)
        tmp.close()
        safe_title = "".join(ch for ch in d.title[:40] if ch.isalnum() or ch in ' _-').strip() or 'document'
        from starlette.background import BackgroundTask
        return FileResponse(tmp.name, filename=f'{d.number} {safe_title}.pdf', media_type='application/pdf',
                            background=BackgroundTask(os.unlink, tmp.name))
    except Exception:
        return JSONResponse(status_code=500, content={"detail": f"PDF error: {traceback.format_exc()[-2000:]}"})


# ============ IMPORT (PDF / DOCX) ============

@app.post("/api/documents/import", response_model=DocumentOut)
async def import_document(
    file: UploadFile = File(...),
    doc_type: str = Form("other"),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    filename = file.filename or "file"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in ("pdf", "docx"):
        raise HTTPException(400, "Поддерживаются только PDF и DOCX файлы")

    content_bytes = await file.read()
    if len(content_bytes) > 20 * 1024 * 1024:
        raise HTTPException(400, "Файл слишком большой (макс. 20 МБ)")

    title = filename.rsplit(".", 1)[0] if "." in filename else filename
    text = ""

    if ext == "docx":
        from docx import Document as DocxDocument
        import io
        try:
            dx = DocxDocument(io.BytesIO(content_bytes))
            text = "\n".join(p.text for p in dx.paragraphs if p.text.strip())
        except Exception:
            raise HTTPException(400, "Не удалось прочитать DOCX файл")

    elif ext == "pdf":
        import tempfile
        try:
            import fitz
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            tmp.write(content_bytes)
            tmp.close()
            pdf_doc = fitz.open(tmp.name)
            pages = []
            for page in pdf_doc:
                pages.append(page.get_text())
            pdf_doc.close()
            os.unlink(tmp.name)
            text = "\n".join(pages)
        except ImportError:
            raise HTTPException(500, "PDF-чтение не поддерживается на сервере")
        except Exception:
            raise HTTPException(400, "Не удалось прочитать PDF файл")

    text = sanitize(text.strip())
    if not text:
        text = "(Содержимое не удалось извлечь)"

    number = gen_number(db, doc_type)
    doc = Document(
        number=number, title=sanitize(title), description=f"Импорт из {filename}",
        content=text, doc_type=doc_type, status="draft",
        priority="normal", extra_fields="{}",
        author_id=user.id,
    )
    db.add(doc)
    db.flush()

    # Save original file as attachment
    doc_dir = os.path.join(UPLOAD_DIR, str(doc.id))
    os.makedirs(doc_dir, exist_ok=True)
    safe_name = secrets.token_hex(8) + "_" + filename
    filepath = os.path.join(doc_dir, safe_name)
    with open(filepath, "wb") as f:
        f.write(content_bytes)
    db.add(Attachment(
        document_id=doc.id, filename=filename,
        filepath=filepath, size=str(len(content_bytes)),
        filesize=len(content_bytes),
    ))

    add_history(db, doc, user.name, f"Импорт из {ext.upper()}: {filename}")
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


# ============ UPLOAD CORRESPONDENCE (PDF) ============

@app.post("/api/documents/upload-correspondence")
async def upload_correspondence(
    file: UploadFile = File(...),
    direction: str = Form("incoming"),
    doc_type: str = Form(""),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Загрузить PDF и создать входящий/исходящий документ с авто-номером."""
    filename = file.filename or "file.pdf"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext != "pdf":
        raise HTTPException(400, "Поддерживаются только PDF файлы")

    content_bytes = await file.read()
    if len(content_bytes) > 20 * 1024 * 1024:
        raise HTTPException(400, "Файл слишком большой (макс. 20 МБ)")

    # Determine doc_type
    if not doc_type:
        doc_type = "incoming_letter" if direction == "incoming" else "outgoing_letter"
    valid_types = INCOMING_TYPES if direction == "incoming" else OUTGOING_TYPES
    if doc_type not in valid_types:
        raise HTTPException(400, f"Недопустимый тип документа для {direction}")

    # Extract text from PDF
    text = ""
    try:
        import fitz, tempfile
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.write(content_bytes)
        tmp.close()
        pdf_doc = fitz.open(tmp.name)
        pages = [page.get_text() for page in pdf_doc]
        pdf_doc.close()
        os.unlink(tmp.name)
        text = "\n".join(pages)
    except ImportError:
        pass
    except Exception:
        pass

    text = sanitize(text.strip()) if text.strip() else "(Содержимое PDF)"
    title = filename.rsplit(".", 1)[0] if "." in filename else filename

    # Create document with auto-number
    number = gen_number(db, doc_type)
    doc = Document(
        number=number, title=sanitize(title), description=f"Загружен PDF: {filename}",
        content=text, doc_type=doc_type, status="draft",
        priority="normal", extra_fields="{}",
        author_id=user.id,
    )
    db.add(doc)
    db.flush()

    # Save file as attachment
    doc_dir = os.path.join(UPLOAD_DIR, str(doc.id))
    os.makedirs(doc_dir, exist_ok=True)
    safe_name = secrets.token_hex(8) + "_" + filename
    filepath = os.path.join(doc_dir, safe_name)
    with open(filepath, "wb") as f:
        f.write(content_bytes)
    db.add(Attachment(
        document_id=doc.id, filename=filename,
        filepath=filepath, size=str(len(content_bytes)),
        filesize=len(content_bytes),
    ))

    dir_label = "входящий" if direction == "incoming" else "исходящий"
    add_history(db, doc, user.name, f"Загружен {dir_label} PDF: {filename}, присвоен №{number}")
    add_audit(db, user.id, user.name, "upload_correspondence", "document", doc.id, f"{dir_label} №{number}")
    db.commit()
    return doc_to_out(load_doc(db, doc.id))


# ============ REGISTRATION JOURNAL (Журнал регистрации) ============

@app.get("/api/journal/{direction}")
def get_registration_journal(
    direction: str,
    date_from: str = "",
    date_to: str = "",
    doc_type: str = "",
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Журнал регистрации входящих/исходящих документов."""
    types = INCOMING_TYPES if direction == "incoming" else OUTGOING_TYPES
    query = db.query(Document).options(
        joinedload(Document.author_user),
        joinedload(Document.approvals).joinedload(Approval.user),
    ).filter(Document.doc_type.in_(types), Document.deleted == False)

    if doc_type and doc_type in types:
        query = query.filter(Document.doc_type == doc_type)
    if date_from:
        try:
            df = datetime.strptime(date_from, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            query = query.filter(Document.created_at >= df)
        except ValueError:
            pass
    if date_to:
        try:
            dt = datetime.strptime(date_to, "%Y-%m-%d").replace(tzinfo=timezone.utc, hour=23, minute=59, second=59)
            query = query.filter(Document.created_at <= dt)
        except ValueError:
            pass

    docs = query.order_by(Document.created_at.asc()).all()
    is_incoming = direction == "incoming"
    entries = []
    for d in docs:
        ef = json.loads(d.extra_fields) if isinstance(d.extra_fields, str) else (d.extra_fields or {})
        counterparty = ef.get("sender") or ef.get("supplier") or ef.get("seller") or ef.get("counterparty") or ef.get("recipient") or ef.get("receiver") or ef.get("buyer") or ""
        reg_num = ef.get("incoming_number") or ef.get("outgoing_number") or ""
        reg_date = ef.get("received_date") or ef.get("send_date") or ""
        amount = ef.get("amount") or ef.get("total_amount") or ""
        approvers = ", ".join(a.user.name for a in d.approvals if a.user) if d.approvals else ""
        status_text = {"draft": "Черновик", "pending": "На согласовании", "approved": "Согласован", "rejected": "Отклонён", "archived": "Архив", "resolved": "Исполнен"}.get(d.status, d.status)
        entries.append({
            "id": d.id,
            "reg_number": d.number or "",
            "ext_number": reg_num,
            "date": str(d.created_at.strftime("%d.%m.%Y %H:%M") if d.created_at else ""),
            "ext_date": reg_date,
            "title": d.title,
            "doc_type": d.doc_type,
            "counterparty": counterparty,
            "author": d.author_user.name if d.author_user else "",
            "amount": str(amount) if amount else "",
            "status": status_text,
            "approvers": approvers,
            "description": (d.description or "")[:100],
        })
    return {"direction": direction, "entries": entries, "total": len(entries)}


# ============ NOMENCLATURE (Номенклатура дел) ============

@app.get("/api/nomenclature")
def list_nomenclature(db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    from sqlalchemy import func
    cases = db.query(NomenclatureCase).order_by(NomenclatureCase.index).all()
    # Count docs per case
    counts = dict(db.query(Document.case_id, func.count(Document.id)).filter(
        Document.case_id.isnot(None), Document.deleted == False
    ).group_by(Document.case_id).all())
    result = []
    for c in cases:
        result.append({
            "id": c.id, "index": c.index, "title": c.title,
            "department": c.department or "", "retention_years": c.retention_years,
            "description": c.description or "",
            "created_at": c.created_at.isoformat() if c.created_at else "",
            "doc_count": counts.get(c.id, 0),
        })
    return result


@app.post("/api/nomenclature")
def create_nomenclature(data: NomenclatureCaseCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только администратор может управлять номенклатурой")
    case = NomenclatureCase(
        index=sanitize(data.index), title=sanitize(data.title),
        department=sanitize(data.department), retention_years=data.retention_years,
        description=sanitize(data.description),
    )
    db.add(case)
    add_audit(db, user.id, user.name, "create", "nomenclature", details=f"{data.index}: {data.title[:40]}")
    db.commit()
    db.refresh(case)
    return {"id": case.id, "index": case.index, "title": case.title,
            "department": case.department, "retention_years": case.retention_years,
            "description": case.description, "created_at": case.created_at.isoformat(), "doc_count": 0}


@app.put("/api/nomenclature/{case_id}")
def update_nomenclature(case_id: int, data: NomenclatureCaseCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только администратор")
    case = db.query(NomenclatureCase).filter(NomenclatureCase.id == case_id).first()
    if not case:
        raise HTTPException(404, "Дело не найдено")
    case.index = sanitize(data.index)
    case.title = sanitize(data.title)
    case.department = sanitize(data.department)
    case.retention_years = data.retention_years
    case.description = sanitize(data.description)
    db.commit()
    return {"ok": True}


@app.delete("/api/nomenclature/{case_id}")
def delete_nomenclature(case_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(403, "Только администратор")
    case = db.query(NomenclatureCase).filter(NomenclatureCase.id == case_id).first()
    if not case:
        raise HTTPException(404, "Дело не найдено")
    # Unlink documents
    db.query(Document).filter(Document.case_id == case_id).update({"case_id": None})
    db.delete(case)
    add_audit(db, user.id, user.name, "delete", "nomenclature", case_id, f"{case.index}: {case.title[:40]}")
    db.commit()
    return {"ok": True}


@app.post("/api/documents/{doc_id}/assign-case")
def assign_doc_to_case(doc_id: int, case_id: int = 0, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    """Привязать документ к делу номенклатуры."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    if case_id:
        case = db.query(NomenclatureCase).filter(NomenclatureCase.id == case_id).first()
        if not case:
            raise HTTPException(404, "Дело не найдено")
    doc.case_id = case_id if case_id else None
    db.commit()
    return {"ok": True}


@app.get("/api/nomenclature/{case_id}/documents")
def get_case_documents(case_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    """Получить документы, привязанные к делу."""
    docs = db.query(Document).filter(Document.case_id == case_id, Document.deleted == False).order_by(Document.created_at.desc()).all()
    return [{"id": d.id, "number": d.number or "", "title": d.title, "doc_type": d.doc_type, "status": d.status, "created_at": str(d.created_at)} for d in docs]


# ============ QR CODE ============

@app.get("/api/documents/{doc_id}/qr")
def get_qr_code(doc_id: int, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    """Generate QR code for document URL as base64 PNG."""
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    try:
        import qrcode, io, base64
        from qrcode.constants import ERROR_CORRECT_M
        qr = qrcode.QRCode(version=1, error_correction=ERROR_CORRECT_M, box_size=6, border=2)
        host = os.getenv("RENDER_EXTERNAL_URL", os.getenv("BASE_URL", ""))
        url = f"{host}/#doc/{doc_id}" if host else f"ЭДО Документ #{doc.number or doc_id}: {doc.title}"
        qr.add_data(url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode()
        return {"qr_base64": b64, "url": url}
    except ImportError:
        raise HTTPException(500, "QR library not installed")


# ============ REPORTS ============

@app.get("/api/reports")
def get_reports(
    date_from: str = "", date_to: str = "",
    doc_type: str = "", status: str = "",
    db: Session = Depends(get_db), user: User = Depends(get_current_user),
):
    """Generate document reports for a date period."""
    q = db.query(Document).filter(Document.deleted == False)
    if user.role != "admin":
        q = q.filter(
            (Document.author_id == user.id) |
            Document.id.in_(db.query(Approval.document_id).filter(Approval.user_id == user.id))
        )
    if date_from:
        try:
            df = datetime.strptime(date_from, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            q = q.filter(Document.created_at >= df)
        except ValueError:
            pass
    if date_to:
        try:
            dt = datetime.strptime(date_to, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            from datetime import timedelta
            q = q.filter(Document.created_at < dt + timedelta(days=1))
        except ValueError:
            pass
    if doc_type:
        q = q.filter(Document.doc_type == doc_type)
    if status:
        q = q.filter(Document.status == status)
    docs = q.order_by(Document.created_at.desc()).all()

    by_status = {}
    by_type = {}
    by_author = {}
    for d in docs:
        by_status[d.status] = by_status.get(d.status, 0) + 1
        by_type[d.doc_type] = by_type.get(d.doc_type, 0) + 1
        author_name = d.author_user.name if d.author_user else "Неизвестен"
        by_author[author_name] = by_author.get(author_name, 0) + 1

    items = [{
        "id": d.id, "number": d.number or "", "title": d.title,
        "doc_type": d.doc_type, "status": d.status, "priority": d.priority or "normal",
        "author_name": d.author_user.name if d.author_user else "",
        "created_at": str(d.created_at), "deadline": d.deadline or "",
    } for d in docs]

    return {
        "total": len(docs),
        "by_status": by_status,
        "by_type": by_type,
        "by_author": by_author,
        "items": items,
        "date_from": date_from,
        "date_to": date_to,
    }


@app.get("/api/reports/export/csv")
def export_reports_csv(
    date_from: str = "", date_to: str = "",
    db: Session = Depends(get_db), user: User = Depends(get_current_user),
):
    """Export reports as CSV file."""
    import csv, io, tempfile
    q = db.query(Document).options(joinedload(Document.author_user)).filter(Document.deleted == False)
    if user.role != "admin":
        q = q.filter(
            (Document.author_id == user.id) |
            Document.id.in_(db.query(Approval.document_id).filter(Approval.user_id == user.id))
        )
    if date_from:
        try:
            df = datetime.strptime(date_from, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            q = q.filter(Document.created_at >= df)
        except ValueError:
            pass
    if date_to:
        try:
            dt = datetime.strptime(date_to, "%Y-%m-%d").replace(tzinfo=timezone.utc)
            from datetime import timedelta
            q = q.filter(Document.created_at < dt + timedelta(days=1))
        except ValueError:
            pass
    docs = q.order_by(Document.created_at.desc()).all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["№", "Номер", "Название", "Тип", "Статус", "Приоритет", "Автор", "Создан", "Дедлайн"])
    for i, d in enumerate(docs, 1):
        writer.writerow([i, d.number or "", d.title, DOC_TYPE_LABELS.get(d.doc_type, d.doc_type),
                         STATUS_LABELS.get(d.status, d.status), d.priority or "",
                         d.author_user.name if d.author_user else "", str(d.created_at)[:19], d.deadline or ""])
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="w", encoding="utf-8-sig")
    tmp.write(output.getvalue())
    tmp.close()
    from starlette.background import BackgroundTask
    return FileResponse(tmp.name, filename="report.csv", media_type="text/csv",
                        background=BackgroundTask(os.unlink, tmp.name))


# ============ ZIP IMPORT ============

@app.post("/api/documents/import-zip")
async def import_zip(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Import multiple documents from a ZIP archive."""
    import zipfile, io
    content_bytes = await file.read()
    if len(content_bytes) > 50 * 1024 * 1024:
        raise HTTPException(400, "Архив слишком большой (макс. 50 МБ)")
    try:
        zf = zipfile.ZipFile(io.BytesIO(content_bytes))
    except zipfile.BadZipFile:
        raise HTTPException(400, "Неверный формат ZIP")

    imported = []
    errors = []
    for name in zf.namelist():
        if name.endswith("/"):
            continue
        ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
        if ext not in ("pdf", "docx"):
            errors.append(f"{name}: неподдерживаемый формат")
            continue
        try:
            file_bytes = zf.read(name)
            if len(file_bytes) > 20 * 1024 * 1024:
                errors.append(f"{name}: файл слишком большой")
                continue
            title = name.rsplit("/", 1)[-1].rsplit(".", 1)[0] if "." in name else name
            text = ""
            if ext == "docx":
                from docx import Document as DocxDocument
                try:
                    dx = DocxDocument(io.BytesIO(file_bytes))
                    text = "\n".join(p.text for p in dx.paragraphs if p.text.strip())
                except Exception:
                    text = "(Не удалось извлечь текст)"
            elif ext == "pdf":
                try:
                    import fitz
                    import tempfile as tf
                    tmp = tf.NamedTemporaryFile(delete=False, suffix=".pdf")
                    tmp.write(file_bytes)
                    tmp.close()
                    pdf_doc = fitz.open(tmp.name)
                    pages = [page.get_text() for page in pdf_doc]
                    pdf_doc.close()
                    os.unlink(tmp.name)
                    text = "\n".join(pages)
                except Exception:
                    text = "(Не удалось извлечь текст)"
            text = sanitize(text.strip()) or "(Содержимое не удалось извлечь)"
            number = gen_number(db, "other")
            doc = Document(
                number=number, title=sanitize(title), description=f"Импорт из ZIP: {name}",
                content=text, doc_type="other", status="draft",
                priority="normal", extra_fields="{}", author_id=user.id,
            )
            db.add(doc)
            db.flush()
            doc_dir = os.path.join(UPLOAD_DIR, str(doc.id))
            os.makedirs(doc_dir, exist_ok=True)
            safe_name = secrets.token_hex(8) + "_" + name.rsplit("/", 1)[-1]
            filepath = os.path.join(doc_dir, safe_name)
            with open(filepath, "wb") as f:
                f.write(file_bytes)
            db.add(Attachment(
                document_id=doc.id, filename=name.rsplit("/", 1)[-1],
                filepath=filepath, size=str(len(file_bytes)), filesize=len(file_bytes),
            ))
            add_history(db, doc, user.name, f"Импорт из ZIP: {name}")
            imported.append({"id": doc.id, "title": title})
        except Exception as e:
            errors.append(f"{name}: {str(e)[:100]}")

    db.commit()
    return {"imported": imported, "errors": errors, "count": len(imported)}


# ============ 1C API ============

API_1C_TOKEN = os.getenv("API_1C_TOKEN", "")


def _check_1c_token(request: Request):
    """Verify 1C API token from header."""
    token = request.headers.get("X-1C-Token", "")
    if not API_1C_TOKEN:
        raise HTTPException(503, "1C API не настроен (задайте API_1C_TOKEN)")
    if token != API_1C_TOKEN:
        raise HTTPException(401, "Неверный токен 1C API")


@app.get("/api/v1/1c/documents")
def api_1c_list_documents(
    request: Request,
    limit: int = 100, offset: int = 0,
    doc_type: str = "", status: str = "",
    date_from: str = "", date_to: str = "",
    db: Session = Depends(get_db),
):
    """1C API: list documents."""
    _check_1c_token(request)
    q = db.query(Document).filter(Document.deleted == False)
    if doc_type:
        q = q.filter(Document.doc_type == doc_type)
    if status:
        q = q.filter(Document.status == status)
    if date_from:
        try:
            q = q.filter(Document.created_at >= datetime.strptime(date_from, "%Y-%m-%d").replace(tzinfo=timezone.utc))
        except ValueError:
            pass
    if date_to:
        try:
            from datetime import timedelta
            q = q.filter(Document.created_at < datetime.strptime(date_to, "%Y-%m-%d").replace(tzinfo=timezone.utc) + timedelta(days=1))
        except ValueError:
            pass
    docs = q.order_by(Document.created_at.desc()).offset(offset).limit(min(limit, 500)).all()
    return [{
        "id": d.id, "number": d.number, "title": d.title, "doc_type": d.doc_type,
        "status": d.status, "priority": d.priority, "deadline": d.deadline or "",
        "author_id": d.author_id, "created_at": d.created_at.isoformat() if d.created_at else "",
        "extra_fields": json.loads(d.extra_fields) if isinstance(d.extra_fields, str) else (d.extra_fields or {}),
    } for d in docs]


@app.get("/api/v1/1c/documents/{doc_id}")
def api_1c_get_document(doc_id: int, request: Request, db: Session = Depends(get_db)):
    """1C API: get single document."""
    _check_1c_token(request)
    doc = db.query(Document).options(
        joinedload(Document.author_user), joinedload(Document.approvals),
    ).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    return {
        "id": doc.id, "number": doc.number, "title": doc.title,
        "description": doc.description, "content": doc.content,
        "doc_type": doc.doc_type, "status": doc.status,
        "priority": doc.priority, "deadline": doc.deadline or "",
        "author_id": doc.author_id,
        "author_name": doc.author_user.name if doc.author_user else "",
        "created_at": doc.created_at.isoformat() if doc.created_at else "",
        "extra_fields": json.loads(doc.extra_fields) if isinstance(doc.extra_fields, str) else (doc.extra_fields or {}),
        "approvals": [{"user_id": a.user_id, "status": a.status, "decided_at": a.decided_at.isoformat() if a.decided_at else ""} for a in doc.approvals],
    }


@app.post("/api/v1/1c/documents")
def api_1c_create_document(request: Request, data: dict, db: Session = Depends(get_db)):
    """1C API: create document."""
    _check_1c_token(request)
    doc_type = data.get("doc_type", "other")
    if doc_type not in VALID_DOC_TYPES:
        doc_type = "other"
    number = gen_number(db, doc_type)
    doc = Document(
        number=number, title=sanitize(data.get("title", "Документ из 1С")),
        description=sanitize(data.get("description", "")),
        content=sanitize(data.get("content", "")),
        doc_type=doc_type, status="draft",
        priority=data.get("priority", "normal"),
        extra_fields=json.dumps(data.get("extra_fields", {}), ensure_ascii=False),
        deadline=data.get("deadline", ""),
        author_id=data.get("author_id", 1),
    )
    db.add(doc)
    db.flush()
    add_history(db, doc, "1С", "Создан через 1С API")
    db.commit()
    return {"id": doc.id, "number": doc.number}


@app.get("/api/v1/1c/users")
def api_1c_list_users(request: Request, db: Session = Depends(get_db)):
    """1C API: list users."""
    _check_1c_token(request)
    users = db.query(User).all()
    return [{"id": u.id, "name": u.name, "department": u.department, "position": u.position, "role": u.role} for u in users]


# ============ WATERMARK & ORG STAMP IN PDF ============

ORG_NAME = os.getenv("ORG_NAME", "ОсОО Эволюшн Групп")
ORG_INN = os.getenv("ORG_INN", "")


# ============ AI ASSISTANT (Groq + Llama) ============

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")

AI_SYSTEM_PROMPT = """Ты — ИИ-помощник системы электронного документооборота (ЭДО).
Ты помогаешь пользователям:
- Суммаризировать документы (краткое содержание)
- Генерировать тексты документов по запросу
- Отвечать на вопросы о документообороте
- Помогать с шаблонами документов
- Давать советы по оформлению
- Навигировать по системе

Отвечай кратко, по делу, на русском языке. Если пользователь просит создать документ — генерируй готовый текст.
Если передан контекст документа — используй его для ответа.

ВАЖНО: Когда пользователь хочет перейти куда-то или создать документ, ОБЯЗАТЕЛЬНО добавь в конец ответа тег навигации в формате [ACTION:страница:параметр].

Доступные действия:
- [ACTION:create:contract] — создать договор
- [ACTION:create:invoice] — создать счёт на оплату
- [ACTION:create:order] — создать приказ
- [ACTION:create:report] — создать отчёт
- [ACTION:create:memo] — создать служебную записку
- [ACTION:create:protocol] — создать протокол
- [ACTION:create:letter] — создать письмо
- [ACTION:create:incoming_letter] — создать входящее письмо
- [ACTION:create:outgoing_letter] — создать исходящее письмо
- [ACTION:create:vacation] — создать заявление на отпуск
- [ACTION:create:trip] — создать командировку
- [ACTION:create:purchase] — создать заявку на закупку
- [ACTION:create:act] — создать акт выполненных работ
- [ACTION:create:statement] — создать заявление
- [ACTION:create:advance_report] — создать авансовый отчёт
- [ACTION:create:payment_order] — создать платёжное поручение
- [ACTION:create:invoice_tax] — создать счёт-фактуру
- [ACTION:create:waybill] — создать товарную накладную
- [ACTION:create:power_of_attorney] — создать доверенность
- [ACTION:create:cash_order] — создать кассовый ордер
- [ACTION:documents] — перейти к списку документов
- [ACTION:approvals] — перейти к согласованиям
- [ACTION:templates] — перейти к шаблонам
- [ACTION:tasks] — перейти к поручениям
- [ACTION:notifications] — перейти к уведомлениям
- [ACTION:archive] — перейти к архиву
- [ACTION:analytics] — перейти к аналитике
- [ACTION:settings] — перейти к настройкам

Примеры (СТРОГО следуй этому формату):
- Пользователь: "Хочу написать заявление на отпуск" → Ты: "Открываю форму создания заявления на отпуск. [ACTION:create:vacation]"
- Пользователь: "Покажи мои документы" → Ты: "Открываю список документов. [ACTION:documents]"
- Пользователь: "Где согласования?" → Ты: "Открываю раздел согласований. [ACTION:approvals]"
- Пользователь: "Создай служебную записку" → Ты: "Открываю форму создания служебной записки. [ACTION:create:memo]"
- Пользователь: "Нужно создать договор" → Ты: "Открываю форму создания договора. [ACTION:create:contract]"
- Пользователь: "Написать входящее письмо" → Ты: "Открываю форму входящего письма. [ACTION:create:incoming_letter]"
- Пользователь: "Покажи шаблоны" → Ты: "Открываю шаблоны документов. [ACTION:templates]"
- Пользователь: "Покажи поручения" → Ты: "Открываю раздел поручений. [ACTION:tasks]"
- Пользователь: "Покажи аналитику" → Ты: "Открываю аналитику. [ACTION:analytics]"
- Пользователь: "Настройки" → Ты: "Открываю настройки. [ACTION:settings]"

ПРАВИЛО: Если пользователь хочет СОЗДАТЬ документ, ПЕРЕЙТИ на страницу, ПОСМОТРЕТЬ раздел — ты ОБЯЗАН добавить тег [ACTION:...] в конец ответа. Это критически важно для навигации. Не забывай тег! Одно короткое предложение + тег.

АВТОЗАПОЛНЕНИЕ ДОКУМЕНТОВ:
Когда пользователь просит создать документ И даёт конкретные данные (даты, суммы, имена, причины), ты ОБЯЗАН добавить тег [FILL:json] с данными для заполнения формы.

Формат: [FILL:{"title":"...","description":"...","content":"...","deadline":"YYYY-MM-DD","priority":"normal"}]

Поля в FILL:
- title — название документа
- description — краткое описание
- content — полный текст документа
- deadline — дедлайн в формате YYYY-MM-DD (если есть)
- priority — приоритет: low, normal, high, urgent

Примеры:
- "Хочу взять отпуск с 15.07.2026 по 25.07.2026" → "Создаю заявление на отпуск. [ACTION:create:vacation] [FILL:{"title":"Заявление на отпуск с 15.07.2026 по 25.07.2026","description":"Ежегодный оплачиваемый отпуск","content":"Прошу предоставить мне ежегодный оплачиваемый отпуск с 15 июля 2026 г. по 25 июля 2026 г. сроком на 11 календарных дней.","deadline":"2026-07-15","priority":"normal"}]"
- "Напиши служебную записку о закупке 5 ноутбуков на сумму 500000 руб" → "Создаю служебную записку. [ACTION:create:memo] [FILL:{"title":"Служебная записка о закупке ноутбуков","description":"Запрос на закупку 5 ноутбуков","content":"Прошу согласовать закупку 5 ноутбуков для отдела разработки на общую сумму 500 000 руб.\\n\\nОбоснование: текущее оборудование устарело и не обеспечивает необходимую производительность.","priority":"normal"}]"
- "Создай приказ о назначении Петрова ответственным за склад" → "Создаю приказ. [ACTION:create:order] [FILL:{"title":"Приказ о назначении ответственного за склад","description":"Назначение Петрова С.А. ответственным за склад","content":"ПРИКАЗ\\n\\nО назначении ответственного за складское хозяйство\\n\\nНазначить Петрова С.А. ответственным за складское хозяйство с 01.08.2026.\\n\\nКонтроль за исполнением приказа оставляю за собой.\\n\\nДиректор ___________","priority":"normal"}]"

ВАЖНО: Генерируй полный текст документа в поле content. Используй все данные которые дал пользователь. JSON в FILL должен быть валидным (экранируй кавычки и переносы строк)."""


def _ai_chat(system_prompt: str, messages: list[dict]) -> str:
    """Send messages to Groq and return response text."""
    if not GROQ_API_KEY:
        raise HTTPException(500, "GROQ_API_KEY не настроен на сервере")
    from groq import Groq
    client = Groq(api_key=GROQ_API_KEY)
    msgs = [{"role": "system", "content": system_prompt}]
    for m in messages:
        role = m.get("role", "user")
        if role in ("user", "assistant"):
            msgs.append({"role": role, "content": m.get("content", "")})
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=msgs,
        max_tokens=2000,
        temperature=0.7,
    )
    return response.choices[0].message.content


@app.post("/api/ai/chat")
def ai_chat(
    request: dict,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    message = request.get("message", "").strip()
    doc_id = request.get("doc_id")
    history = request.get("history", [])

    if not message:
        raise HTTPException(400, "Сообщение не может быть пустым")

    doc_context = ""
    if doc_id:
        doc = db.query(Document).filter(Document.id == doc_id).first()
        if doc:
            check_doc_access(doc, user, db)
            doc_context = f"\n\nКонтекст документа:\nНазвание: {doc.title}\nТип: {DOC_TYPE_LABELS.get(doc.doc_type, doc.doc_type)}\nСтатус: {STATUS_LABELS.get(doc.status, doc.status)}\nОписание: {doc.description}\nСодержание: {doc.content[:3000]}"

    system = AI_SYSTEM_PROMPT + doc_context

    msgs = []
    for h in history[-10:]:
        role = h.get("role", "user")
        if role in ("user", "assistant"):
            msgs.append({"role": role, "content": h.get("content", "")})
    msgs.append({"role": "user", "content": message})

    try:
        reply = _ai_chat(system, msgs)
        # Parse ACTION tags
        action = None
        action_match = re.search(r'\[ACTION:([^\]]+)\]', reply)
        if action_match:
            parts = action_match.group(1).split(":")
            reply = re.sub(r'\s*\[ACTION:[^\]]+\]', '', reply).strip()
            action = {"page": parts[0]}
            if len(parts) > 1:
                action["param"] = parts[1]
        # Parse FILL tags
        fill = None
        fill_match = re.search(r'\[FILL:(\{.*?\})\]', reply, re.DOTALL)
        if fill_match:
            try:
                fill = json.loads(fill_match.group(1))
            except json.JSONDecodeError:
                pass
            reply = re.sub(r'\s*\[FILL:\{.*?\}\]', '', reply, flags=re.DOTALL).strip()
        return {"reply": reply, "action": action, "fill": fill}
    except Exception as e:
        raise HTTPException(500, f"Ошибка ИИ: {str(e)}")


@app.post("/api/ai/summarize")
def ai_summarize(
    request: dict,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    doc_id = request.get("doc_id")
    if not doc_id:
        raise HTTPException(400, "doc_id обязателен")

    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(404, "Документ не найден")
    check_doc_access(doc, user, db)

    text = f"Название: {doc.title}\nТип: {DOC_TYPE_LABELS.get(doc.doc_type, doc.doc_type)}\nОписание: {doc.description}\nСодержание:\n{doc.content[:4000]}"

    try:
        reply = _ai_chat(
            "Ты суммаризатор документов. Дай краткое содержание документа в 3-5 предложениях на русском языке.",
            [{"role": "user", "content": text}],
        )
        return {"summary": reply}
    except Exception as e:
        raise HTTPException(500, f"Ошибка ИИ: {str(e)}")


@app.post("/api/ai/generate")
def ai_generate(
    request: dict,
    db: Session = Depends(get_db),
    user: User = Depends(get_current_user),
):
    prompt = request.get("prompt", "").strip()
    doc_type = request.get("doc_type", "")

    if not prompt:
        raise HTTPException(400, "Опишите что нужно сгенерировать")

    type_label = DOC_TYPE_LABELS.get(doc_type, "документ")
    system = f'Ты генератор документов для системы ЭДО. Сгенерируй текст документа типа "{type_label}" по запросу пользователя. Формат: готовый текст документа, который можно сразу использовать. Без лишних пояснений.'

    try:
        reply = _ai_chat(system, [{"role": "user", "content": prompt}])
        return {"content": reply}
    except Exception as e:
        raise HTTPException(500, f"Ошибка ИИ: {str(e)}")


app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/")
def index():
    return FileResponse("static/index.html")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=int(os.getenv("PORT", 8000)), reload=True)
